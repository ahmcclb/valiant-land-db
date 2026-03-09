from flask import Flask, render_template, render_template_string, jsonify, request, send_file, send_from_directory
import psycopg2
from psycopg2.extras import RealDictCursor
import os
import io
import csv
import json
import openpyxl
from openpyxl.utils import get_column_letter
from copy import deepcopy
from datetime import datetime, timedelta
from database_pg import get_db_connection
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re
import logging
import sys

logger = logging.getLogger(__name__)

if getattr(sys, 'frozen', False):
    # EXE mode - log to file next to EXE
    log_path = os.path.join(os.path.dirname(sys.executable), 'app_debug.log')
    logging.basicConfig(
        filename=log_path,
        level=logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    # Also add console handler for pywebview
    console = logging.StreamHandler()
    console.setLevel(logging.DEBUG)
    logging.getLogger('').addHandler(console)

# === STEP 1: Define directories for bundled assets vs persistent data ===
def get_bundle_dir():
    """Get directory containing bundled static assets (CSS, JS, templates, logos).
    In PyInstaller onefile mode: temp extraction folder (sys._MEIPASS).
    In development: script folder."""
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        return sys._MEIPASS
    else:
        return os.path.dirname(os.path.abspath(__file__))

def get_base_dir():
    """Get the base directory for persistent data (uploads, config, exports).
    This is always the folder containing the EXE or script."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

BUNDLE_DIR = get_bundle_dir()   # For CSS, JS, templates (PyInstaller temp when frozen)
BASE_DIR = get_base_dir()       # For uploads, exports, config.json (persistent)

# === STEP 2: Initialize Flask WITHOUT automatic static handling ===
# We will manually handle /static/ to check both persistent and bundled folders
app = Flask(__name__, 
            static_folder=None,  # CRITICAL: Disables Flask's default static route
            template_folder=os.path.join(BUNDLE_DIR, 'templates'))

# === STEP 3: Define persistent data paths using BASE_DIR ===
STATIC_PATH = os.path.join(BASE_DIR, 'static')  # User uploads (persistent)
TEMPLATE_PATH = os.path.join(BUNDLE_DIR, 'templates', 'documents')
GENERATED_PATH = os.path.join(STATIC_PATH, 'generated_documents')
TEMP_PATH = os.path.join(STATIC_PATH, 'temp_documents')

os.makedirs(GENERATED_PATH, exist_ok=True)
os.makedirs(TEMP_PATH, exist_ok=True)

# === STEP 4: Ensure upload directories exist in PERSISTENT location ===
os.makedirs(os.path.join(STATIC_PATH, 'uploads', 'photos'), exist_ok=True)
os.makedirs(os.path.join(STATIC_PATH, 'uploads', 'documents'), exist_ok=True)
os.makedirs(os.path.join(STATIC_PATH, 'uploads', 'company'), exist_ok=True)

# Document generation progress storage (in production use Redis)
generation_progress = {}

@app.route('/api/sync', methods=['POST'])
def api_sync():
    """Execute sync with cloud"""
    from sync_service import ValiantLandSync
    import os
    
    data = request.json
    direction = data.get('direction', 'bidirectional')
    
    # Config from environment variables or config file
    supabase_url = os.getenv('SUPABASE_URL', 'https://gkgiuokglsgsywctpohe.supabase.co')
    supabase_key = os.getenv('SUPABASE_KEY', 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImdrZ2l1b2tnbHNnc3l3Y3Rwb2hlIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc3MTI3NDQ0MywiZXhwIjoyMDg2ODUwNDQzfQ.UY67-jQ2Xzrh0EWxtH6Jx5_Ff2ygqBgAXIVay9XVfsw')
    
    local_config = {
        'host': 'localhost',
        'database': 'valiant_land',
        'user': 'vl_user',
        'password': 'ifoLdouTEAThleCt'
    }
    
    try:
        sync = ValiantLandSync(local_config, supabase_url, supabase_key)
        stats = sync.sync_database(direction)
        
        # Also sync files if requested
        if data.get('include_files', True):
            file_stats = sync.sync_files(direction)  # Pass the direction!
            stats.update(file_stats)
        
        return jsonify({
            'success': True,
            **stats
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/sync/status', methods=['GET'])
def sync_status():
    """Get current sync status"""
    try:
        from sync_service import ValiantLandSync
        # Pass config explicitly or let it load from config.json
        sync = ValiantLandSync()
        status = sync.get_sync_status()
        return jsonify(status)
    except Exception as e:
        logger.error(f"Error getting sync status: {e}")
        return jsonify({
            'error': str(e),
            'properties_pending': 0,
            'properties_synced': 0,
            'properties_conflicts': 0,
            'owners_pending': 0,
            'owners_synced': 0,
            'files_pending': 0,
            'last_sync': None
        }), 500

@app.route('/favicon.ico')
def favicon():
    """Serve favicon.ico from bundled static folder"""
    return send_from_directory(
        os.path.join(BUNDLE_DIR, 'static'),
        'favicon.ico',
        mimetype='image/x-icon'
    )

# State abbreviation to full name mapping
STATE_ABBR_TO_FULL = {
    'AL': 'Alabama', 'AK': 'Alaska', 'AZ': 'Arizona', 'AR': 'Arkansas', 'CA': 'California',
    'CO': 'Colorado', 'CT': 'Connecticut', 'DE': 'Delaware', 'DC': 'District of Columbia',
    'FL': 'Florida', 'GA': 'Georgia', 'HI': 'Hawaii', 'ID': 'Idaho', 'IL': 'Illinois',
    'IN': 'Indiana', 'IA': 'Iowa', 'KS': 'Kansas', 'KY': 'Kentucky', 'LA': 'Louisiana',
    'ME': 'Maine', 'MD': 'Maryland', 'MA': 'Massachusetts', 'MI': 'Michigan', 'MN': 'Minnesota',
    'MS': 'Mississippi', 'MO': 'Missouri', 'MT': 'Montana', 'NE': 'Nebraska', 'NV': 'Nevada',
    'NH': 'New Hampshire', 'NJ': 'New Jersey', 'NM': 'New Mexico', 'NY': 'New York',
    'NC': 'North Carolina', 'ND': 'North Dakota', 'OH': 'Ohio', 'OK': 'Oklahoma', 'OR': 'Oregon',
    'PA': 'Pennsylvania', 'RI': 'Rhode Island', 'SC': 'South Carolina', 'SD': 'South Dakota',
    'TN': 'Tennessee', 'TX': 'Texas', 'UT': 'Utah', 'VT': 'Vermont', 'VA': 'Virginia',
    'WA': 'Washington', 'WV': 'West Virginia', 'WI': 'Wisconsin', 'WY': 'Wyoming',
    'AS': 'American Samoa', 'GU': 'Guam', 'MP': 'Northern Mariana Islands', 'PR': 'Puerto Rico',
    'VI': 'U.S. Virgin Islands', 'UM': 'U.S. Minor Outlying Islands',
    'AE': 'Armed Forces Europe', 'AP': 'Armed Forces Pacific', 'AA': 'Armed Forces Americas',
    'AB': 'Alberta', 'BC': 'British Columbia', 'MB': 'Manitoba', 'NB': 'New Brunswick',
    'NL': 'Newfoundland and Labrador', 'NS': 'Nova Scotia', 'ON': 'Ontario',
    'PE': 'Prince Edward Island', 'QC': 'Quebec', 'SK': 'Saskatchewan',
    'NT': 'Northwest Territories', 'YT': 'Yukon', 'NU': 'Nunavut'
}

def get_template_path(filename):
    """Check external folder (next to EXE) first, then bundled"""
    if getattr(sys, 'frozen', False):
        # When frozen, check next to EXE first (editable templates)
        external = os.path.join(os.path.dirname(sys.executable), 'static', 'document_templates', filename)
        if os.path.exists(external):
            return external
        # Fall back to bundled template inside EXE
        return os.path.join(TEMPLATE_PATH, filename)
    else:
        # Development mode - use local templates folder
        return os.path.join(TEMPLATE_PATH, filename)

def get_longstate_from_abbr(abbr):
    """Convert state abbreviation to full name"""
    if not abbr:
        return ''
    return STATE_ABBR_TO_FULL.get(abbr.upper(), abbr)

def clean_numeric_field(value):
    """Convert empty strings to None for numeric fields."""
    if value is None:
        return None
    if isinstance(value, str) and value.strip() == '':
        return None
    return value

def clean_boolean_field(value):
    """Convert empty strings and various inputs to proper boolean/None."""
    if value is None or value == '':
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.lower() in ('yes', 'true', '1', 'on', 'y')
    return bool(value)

def mark_record_modified(table, record_id, conn=None):
    """Mark a record as pending sync whenever it's modified"""
    should_close = False
    if conn is None:
        conn = get_db_connection()
        should_close = True
    
    try:
        cursor = conn.cursor()
        id_column = 'p_id' if table == 'properties' else 'or_id'
        cursor.execute(f"""
            UPDATE {table} 
            SET sync_status = 'pending', 
                modified_at = NOW(),
                sync_version = sync_version + 1
            WHERE {id_column} = %s
        """, (record_id,))
        
        if should_close:
            conn.commit()
            conn.close()
    except Exception as e:
        print(f"Error marking {table} {record_id} as modified: {e}")

@app.route('/')
def dashboard():
    """Main Dashboard Page."""
    return render_template('dashboard.html')

@app.route('/new-offer-request')
def new_offer_request():
    """New Offer Request page - creates property records from caller conversations."""
    return render_template('new_offer_request.html')

@app.route('/api/dashboard/properties')
def get_properties():
    """API endpoint to fetch properties for the table with filtering, pagination, and column selection."""
    
    # Get query parameters
    filter_type = request.args.get('filter', 'all_in_process')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 10))
    search = request.args.get('search', '')
    sort_by = request.args.get('sort_by', 'p_id')
    sort_dir = request.args.get('sort_dir', 'asc')
    
    # Get visible columns from parameter (comma-separated list of field names)
    visible_columns = request.args.get('columns', 
        'p_id,p_status,p_apn,or_name,p_county,p_state,p_acres,p_comp_market_value')
    visible_columns = visible_columns.split(',')
    
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    # Build base query with JOIN to owners table
    base_query = """
        FROM properties p
        JOIN owners o ON p.or_id = o.or_id
        LEFT JOIN statuses s ON p.p_status_id = s.status_id
        WHERE 1=1
    """
    params = []
    
    # Apply filter logic
    if filter_type == 'all_in_process':
        # Exclude: Sold, On Hold, Skip Trace, Closed
        cursor.execute("""
            SELECT status_id FROM statuses 
            WHERE s_status IN ('SOLD', 'ON HOLD', 'Skip Trace', 'FILE CLOSED')
        """)
        excluded_ids = [row['status_id'] for row in cursor.fetchall()]
        if excluded_ids:
            base_query += f" AND p.p_status_id NOT IN ({','.join(['%s'] * len(excluded_ids))})"
            params.extend(excluded_ids)
            
    elif filter_type == 'prospects':
        cursor.execute("""
            SELECT status_id FROM statuses 
            WHERE s_status IN ('Prospect', 'Mailed Letter 1')
        """)
        status_ids = [row['status_id'] for row in cursor.fetchall()]
        if status_ids:
            base_query += f" AND p.p_status_id IN ({','.join(['%s'] * len(status_ids))})"
            params.extend(status_ids)
            
    elif filter_type == 'skip_trace':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'SKIP TRACE'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'offer_requests':
        cursor.execute("""
            SELECT status_id FROM statuses 
            WHERE s_status IN ('Pending Preliminary Research', 'Ready to Send Offer', 'Ready to Email Offer')
        """)
        status_ids = [row['status_id'] for row in cursor.fetchall()]
        if status_ids:
            base_query += f" AND p.p_status_id IN ({','.join(['%s'] * len(status_ids))})"
            params.extend(status_ids)
            
    elif filter_type == 'offers_made':
        cursor.execute("""
            SELECT status_id FROM statuses 
            WHERE s_status IN ('Offers Sent', 'Offers Emailed')
        """)
        status_ids = [row['status_id'] for row in cursor.fetchall()]
        if status_ids:
            base_query += f" AND p.p_status_id IN ({','.join(['%s'] * len(status_ids))})"
            params.extend(status_ids)
            
    elif filter_type == 'second_offers':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = '2nd Offer Sent'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'buying':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'Open Escrow - Detailed Research'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'selling':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'Complete/ Ready To Sell'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'sold':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'SOLD'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'on_hold':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'ON HOLD'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])
            
    elif filter_type == 'all_closed':
        cursor.execute("SELECT status_id FROM statuses WHERE s_status = 'FILE CLOSED'")
        row = cursor.fetchone()
        if row:
            base_query += " AND p.p_status_id = %s"
            params.append(row['status_id'])

    elif filter_type == 'expired':
        # Contract expired: p_contract_expires_date is MM/DD/YYYY text
            base_query += """ AND (p_contract_expires_date IS NOT NULL 
                           AND p_contract_expires_date != '' 
                           AND TO_DATE(p_contract_expires_date, 'MM/DD/YYYY') < CURRENT_DATE)"""
                           
    elif filter_type == 'date_created':
        # Date created range (format: "Dec 18, 2018, 7:38 pm")
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        if date_from:
            base_query += " AND TO_TIMESTAMP(p.p_create_time, 'Mon DD, YYYY, HH12:MI pm') >= TO_TIMESTAMP(%s, 'Mon DD, YYYY, HH12:MI pm')"
            params.append(date_from)
        if date_to:
            base_query += " AND TO_TIMESTAMP(p.p_create_time, 'Mon DD, YYYY, HH12:MI pm') <= TO_TIMESTAMP(%s, 'Mon DD, YYYY, HH12:MI pm')"
            params.append(date_to)
            
    # Apply search filter
    if search:
        base_query += """ AND (
            p.p_apn LIKE %s OR 
            o.or_fname || ' ' || o.or_lname LIKE %s OR 
            o.o_company LIKE %s OR 
            p.p_county LIKE %s OR 
            p.p_state LIKE %s
        )"""
        search_term = f'%{search}%'
        params.extend([search_term, search_term, search_term, search_term, search_term])
    
    # NEW: Add county filter if provided
    county_filter = request.args.get('county', '')
    if county_filter:
        base_query += " AND p.p_county LIKE %s"
        params.append(f'%{county_filter}%')
    
    # Count total records
    count_query = f"SELECT COUNT(*) {base_query}"
    cursor.execute(count_query, params)
    total_records = cursor.fetchone()['count']
    total_pages = (total_records + per_page - 1) // per_page
    
    # Build SELECT clause with visible columns
    select_fields = []
    
    # Always include p_id and or_id for internal use
    select_fields.append('p.p_id')
    select_fields.append('p.or_id')
    
    for col in visible_columns:
        if col == 'p_id' or col == 'or_id':
            continue
        elif col == 'or_name':
            # Special case: calculated name field - uses caller name for individuals
            select_fields.append(
                "CASE WHEN o.o_type = 'Company' THEN COALESCE(o.o_company, '') ELSE TRIM(COALESCE(o.or_fname, '') || ' ' || COALESCE(o.or_lname, '')) END as or_name"
            )
        elif col == 'tags':
            # Special case: get concatenated tags
            select_fields.append("""
                (SELECT STRING_AGG(t.tag_name, ' | ') 
                 FROM property_tags pt 
                 JOIN tags t ON pt.tag_id = t.tag_id 
                 WHERE pt.p_id = p.p_id) as tags
            """)
        elif col == 'p_status':
            select_fields.append('s.s_status as p_status')
        elif col.startswith('or_'):
            select_fields.append(f'o.{col}')
        elif col.startswith('o_'):
            select_fields.append(f'o.{col}')
        else:
            select_fields.append(f'p.{col}')
    
    # Build main query
    fields_str = ', '.join(select_fields)
    query = f"""
        SELECT {fields_str}
        {base_query}
        ORDER BY {sort_by} {sort_dir}
        LIMIT %s OFFSET %s
    """
    params.extend([per_page, (page - 1) * per_page])
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    
    # Convert rows to list of dicts
    properties = []
    for row in rows:
        properties.append(dict(row))
    
    conn.close()

    if properties:
        print(f"DEBUG: First record or_name='{properties[0].get('or_name', 'MISSING')}', or_id={properties[0].get('or_id')}")
    
    return jsonify({
        'properties': properties,
        'pagination': {
            'current_page': page,
            'per_page': per_page,
            'total_records': total_records,
            'total_pages': total_pages
        }
    })

@app.route('/api/dashboard/columns')
def get_available_columns():
    """Return list of all available columns for customization."""
    
    columns = [
        {'field': 'p_id', 'label': 'ID', 'default_visible': True},
        {'field': 'p_apn', 'label': 'APN', 'default_visible': True},
        {'field': 'p_status', 'label': 'Status', 'default_visible': True},
        {'field': 'or_name', 'label': 'Name', 'default_visible': True},
        {'field': 'o_company', 'label': 'Company', 'default_visible': False},
        {'field': 'or_m_address', 'label': 'Mailing Address', 'default_visible': False},
        {'field': 'or_m_address2', 'label': 'Mailing Address 2', 'default_visible': False},
        {'field': 'or_m_city', 'label': 'Mailing City', 'default_visible': False},
        {'field': 'or_m_state', 'label': 'Mailing State', 'default_visible': False},
        {'field': 'or_m_zip', 'label': 'Mailing Zip', 'default_visible': False},
        {'field': 'p_county', 'label': 'Property County', 'default_visible': True},
        {'field': 'p_state', 'label': 'Property State', 'default_visible': True},
        {'field': 'p_acres', 'label': 'Property Size (Acres)', 'default_visible': True},
        {'field': 'p_comp_market_value', 'label': 'Comp Market Value', 'default_visible': True},
        {'field': 'p_city', 'label': 'Property City', 'default_visible': False},
        {'field': 'p_address', 'label': 'Property Address', 'default_visible': False},
        {'field': 'p_base_tax', 'label': 'Annual Base Taxes', 'default_visible': False},
        {'field': 'p_back_tax', 'label': 'Back Taxes', 'default_visible': False},
        {'field': 'p_betty_score', 'label': 'Betty Score', 'default_visible': False},
        {'field': 'p_contract_expires_date', 'label': 'Contract Expires On', 'default_visible': False},
        {'field': 'p_county_assessed_value', 'label': 'County Assessed Value', 'default_visible': False},
        {'field': 'p_county_market_value', 'label': 'County Market Value', 'default_visible': False},
        {'field': 'p_create_time', 'label': 'Create Time', 'default_visible': False},
        {'field': 'p_sale_price', 'label': 'For Sale Price', 'default_visible': False},        
        {'field': 'p_hoa', 'label': 'HOA Fee', 'default_visible': False},
        {'field': 'p_aquired', 'label': 'How Acquired', 'default_visible': False},
        {'field': 'p_last_updated', 'label': 'Last Update', 'default_visible': False},
        {'field': 'p_last_sold_date', 'label': 'Last Sold Date', 'default_visible': False},
        {'field': 'p_last_sold_amount', 'label': 'Last Sold Amount', 'default_visible': False},
        {'field': 'p_last_transaction_date', 'label': 'Last Transaction Date', 'default_visible': False},
        {'field': 'p_last_transaction_doc_type', 'label': 'Last Transaction Doc Type', 'default_visible': False},
        {'field': 'p_owned', 'label': 'Length of Ownership', 'default_visible': False},
        {'field': 'p_liens', 'label': 'Liens', 'default_visible': False},
        {'field': 'p_min_acceptable_offer', 'label': 'Min Acceptable Offer', 'default_visible': False},
        {'field': 'p_note', 'label': 'Notes', 'default_visible': False},
        {'field': 'p_offer_accept_date', 'label': 'Offer Accept By', 'default_visible': False},
        {'field': 'p_price', 'label': 'Offer Amount', 'default_visible': False},
        {'field': 'p_m_date', 'label': 'Offer Mail Date', 'default_visible': False},
        {'field': 'or_id', 'label': 'Owner\'s ID', 'default_visible': False},
        {'field': 'or_phone', 'label': 'Phone', 'default_visible': False},
        {'field': 'p_sqft', 'label': 'Property Size (SQ. FT)', 'default_visible': False},
        {'field': 'p_zip', 'label': 'Property Zip', 'default_visible': False},
        {'field': 'p_purchase_amount', 'label': 'Purchase Amount', 'default_visible': False},
        {'field': 'p_purchase_closing_costs', 'label': 'Purchase Closing Cost', 'default_visible': False},
        {'field': 'p_purchased_on', 'label': 'Purchase Date', 'default_visible': False},
        {'field': 'p_agent_name', 'label': 'Realtor Name', 'default_visible': False},
        {'field': 'p_agent_phone', 'label': 'Realtor Phone', 'default_visible': False},
        {'field': 'p_sold_amount', 'label': 'Sale Amount', 'default_visible': False},
        {'field': 'p_sold_closing_costs', 'label': 'Sale Closing Costs', 'default_visible': False},
        {'field': 'p_sold_on', 'label': 'Sale Date On', 'default_visible': False},
        {'field': 'p_short_legal', 'label': 'Short Legal Description', 'default_visible': False},
        {'field': 'p_status_last_updated', 'label': 'Status Last Update', 'default_visible': False},
        {'field': 'tags', 'label': 'Tags', 'default_visible': False},
        {'field': 'o_type', 'label': 'Type', 'default_visible': False},
        {'field': 'p_viable', 'label': 'Viable Seller', 'default_visible': False},
        {'field': 'p_zoning', 'label': 'Zoning', 'default_visible': False}
    ]
    
    return jsonify({'columns': columns})

@app.route('/api/dashboard/export/csv', methods=['POST'])
def export_csv():
    """Export all data to CSV and return as direct download."""
    from io import StringIO, BytesIO
    
    data = request.json
    records = data.get('records', [])
    
    if not records:
        return jsonify({'error': 'No records to export'}), 400
    
    # Generate filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"Export_{timestamp}.csv"
    
    # Get all field names from first record
    fieldnames = list(records[0].keys()) if records else []
    
    # Create CSV in memory
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for record in records:
        writer.writerow(record)
    
    # Convert to BytesIO for send_file
    bytes_output = BytesIO(output.getvalue().encode('utf-8-sig'))
    bytes_output.seek(0)
    
    return send_file(
        bytes_output,
        mimetype='text/csv',
        as_attachment=True,
        download_name=filename
    )

@app.route('/api/dashboard/export/mailing', methods=['POST'])
def export_mailing():
    """Export mailing list CSV and return as direct download."""
    from io import StringIO, BytesIO
    
    data = request.json
    records = data.get('records', [])
    export_type = data.get('type', 'usmail')
    
    if not records:
        return jsonify({'error': 'No records to export'}), 400
    
    # Generate m_date (today's date in "Month D, YYYY" format)
    today = datetime.now()
    m_date = '="' + today.strftime('%B ') + str(today.day) + today.strftime(', %Y') + '"'
    
    # Deduplicate by or_id, keeping the one with lowest p_id
    seen_or_ids = set()
    deduped_records = []
    for record in sorted(records, key=lambda x: x.get('p_id', 0)):
        or_id = record.get('or_id')
        if or_id not in seen_or_ids:
            seen_or_ids.add(or_id)
            deduped_records.append(record)
    
    # Export all records (no status filtering)
    filtered_records = deduped_records
    
    if not filtered_records:
        return jsonify({'error': 'No records to export'}), 400
    
    # Determine fields to export
    if export_type == 'usmail':
        fields = ['or_id', 'p_state', 'p_county', 'p_apn', 'or_greeting', 
                  'or_name', 'or_m_address', 'or_m_city', 'or_m_state', 
                  'or_m_zip', 'p_longstate', 'm_date']
    else:
        fields = ['or_id', 'p_state', 'p_county', 'p_apn', 'or_greeting', 
                  'or_name', 'or_email', 'or_m_address', 'or_m_city', 
                  'or_m_state', 'or_m_zip', 'p_longstate', 'm_date']
    
    # Generate filename
    if export_type == 'usmail':
        filename = f"Mailing_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    else:
        filename = f"Emailing_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    
    # Prepare data for export
    export_data = []
    for record in filtered_records:
        row = {}
        # Calculate or_greeting and or_name on-the-fly
        o_type = record.get('o_type', '')
        or_fname = record.get('or_fname', '')
        or_lname = record.get('or_lname', '')
        o_company = record.get('o_company', '')
        
        if o_type == 'Company':
            or_name = o_company
            or_greeting = "To whom it may concern,"
        else:
            or_name = f"{or_fname} {or_lname}".strip()
            or_greeting = f"Dear {or_fname},"
        
        for field in fields:
            if field == 'or_greeting':
                row[field] = or_greeting
            elif field == 'or_name':
                row[field] = or_name
            elif field == 'm_date':
                row[field] = m_date
            elif field == 'p_longstate':
                longstate = record.get('p_longstate')
                if not longstate:
                    longstate = get_longstate_from_abbr(record.get('p_state', ''))
                row[field] = longstate
            elif field.startswith('or_m_'):
                row[field] = record.get(field, '')
            elif field in record:
                row[field] = record[field]
            else:
                row[field] = ''
        
        export_data.append(row)
    
    # Create CSV in memory
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=fields)
    writer.writeheader()
    writer.writerows(export_data)
    
    # Convert to BytesIO for send_file
    bytes_output = BytesIO(output.getvalue().encode('utf-8-sig'))
    bytes_output.seek(0)
    
    return send_file(
        bytes_output,
        mimetype='text/csv',
        as_attachment=True,
        download_name=filename
    )

@app.route('/status-management')
def status_management():
    """Status Management Page."""
    return render_template('status_management.html')

@app.route('/api/statuses', methods=['GET', 'POST'])
def handle_statuses():
    """Get all statuses or create a new one."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    if request.method == 'GET':
        cursor.execute('SELECT * FROM statuses ORDER BY s_order ASC')
        statuses = [dict(row) for row in cursor.fetchall()]
        conn.close()
        return jsonify({'statuses': statuses})
    
    elif request.method == 'POST':
        data = request.json
        s_status = data.get('s_status')
        s_color = data.get('s_color', '#808080')
        s_order = data.get('s_order', 0)
        
        cursor.execute('''
            INSERT INTO statuses (s_status, s_color, s_order, s_is_active)
            VALUES (%s, %s, %s, TRUE)
            RETURNING status_id
        ''', (s_status, s_color, s_order))
        
        status_id = cursor.fetchone()['status_id']
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'status_id': status_id})

@app.route('/api/statuses/reorder', methods=['POST'])
def reorder_statuses():
    """Update the order of statuses."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    data = request.json
    statuses = data.get('statuses', [])
    
    for status in statuses:
        cursor.execute('''
            UPDATE statuses SET s_order = %s WHERE status_id = %s
        ''', (status['s_order'], status['status_id']))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/properties/change-status', methods=['POST'])
def change_property_status():
    """Change status for selected property records."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    data = request.json
    p_ids = data.get('p_ids', [])
    new_status_id = data.get('status_id')
    
    if not p_ids or not new_status_id:
        return jsonify({'error': 'Missing required parameters'}), 400
    
    # Get the status name for verification
    cursor.execute('SELECT s_status FROM statuses WHERE status_id = %s', (new_status_id,))
    status_row = cursor.fetchone()
    if not status_row:
        return jsonify({'error': 'Invalid status_id'}), 400
    
    # Update status and timestamp for each property
    from datetime import datetime
    current_timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')
    
    placeholders = ','.join(['%s'] * len(p_ids))
    cursor.execute(f'''
        UPDATE properties 
        SET p_status_id = %s, p_status_last_updated = %s
        WHERE p_id IN ({placeholders})
    ''', [new_status_id, current_timestamp] + p_ids)
    
    conn.commit()
    
    # Mark properties for sync
    for p_id in p_ids:
        mark_record_modified('properties', p_id, conn)    
    
    conn.close()
    
    return jsonify({'success': True, 'updated': len(p_ids)})

@app.route('/api/statuses/<int:status_id>', methods=['DELETE'])
def delete_status(status_id):
    """Delete a status (only if not in use)."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Check if status is in use
    cursor.execute('SELECT COUNT(*) FROM properties WHERE p_status_id = %s', (status_id,))
    count = cursor.fetchone()[0]
    
    if count > 0:
        conn.close()
        return jsonify({'error': 'Cannot delete status that is in use'}), 400
    
    cursor.execute('DELETE FROM statuses WHERE status_id = %s', (status_id,))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/property/<int:p_id>', methods=['DELETE'])
def delete_property(p_id):
    """Delete a property record and propagate to cloud."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 1. Insert tombstone record for sync propagation
        cursor.execute('''
            INSERT INTO sync_deletions (table_name, record_id, sync_status)
            VALUES (%s, %s, %s)
            ON CONFLICT (table_name, record_id) DO UPDATE SET
            deleted_at = NOW(), sync_status = 'pending', cloud_deleted = FALSE
        ''', ('properties', p_id, 'pending'))
        
        # 2. Try to delete from Supabase immediately (if online)
        try:
            from sync_service import ValiantLandSync
            sync = ValiantLandSync()
            sync.supabase.table('properties').delete().eq('p_id', p_id).execute()
            cursor.execute('UPDATE sync_deletions SET cloud_deleted = TRUE WHERE table_name = %s AND record_id = %s', ('properties', p_id))
        except Exception as cloud_error:
            print(f"Cloud delete queued for later (property {p_id}): {cloud_error}")
        
        # 3. Delete related files tracking
        cursor.execute('DELETE FROM file_sync WHERE local_path LIKE %s', (f'%/p_{p_id}/%',))
        
        # 4. Delete local property (cascade will handle photos/docs/links)
        cursor.execute('DELETE FROM properties WHERE p_id = %s', (p_id,))
        
        conn.commit()
        return jsonify({'success': True, 'message': 'Record deleted successfully and queued for cloud sync'})
        
    except Exception as e:
        conn.rollback()
        print(f"ERROR deleting property: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/statuses/<int:status_id>', methods=['PUT'])
def update_status(status_id):
    """Update an existing status."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    data = request.json
    s_status = data.get('s_status')
    s_color = data.get('s_color', '#808080')
    
    if not s_status:
        return jsonify({'error': 'Status name is required'}), 400
    
    cursor.execute('''
        UPDATE statuses 
        SET s_status = %s, s_color = %s
        WHERE status_id = %s
    ''', (s_status, s_color, status_id))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

# TAG MANAGEMENT ROUTES
@app.route('/tags')
def tag_management():
    """Tag Management Page."""
    return render_template('tag_management.html')

@app.route('/api/tags', methods=['GET', 'POST'])
def handle_tags():
    """Get all tags or create a new one."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    if request.method == 'GET':
        cursor.execute('SELECT * FROM tags ORDER BY tag_id ASC')
        tags = [dict(row) for row in cursor.fetchall()]
        conn.close()
        return jsonify({'tags': tags})
    
    elif request.method == 'POST':
        data = request.json
        tag_name = data.get('tag_name')
        
        if not tag_name:
            return jsonify({'error': 'Tag name is required'}), 400
        
        cursor.execute('''
            INSERT INTO tags (tag_name, tag_color)
            VALUES (%s, '#808080')
            RETURNING tag_id
        ''', (tag_name,))
        
        tag_id = cursor.fetchone()['tag_id']
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'tag_id': tag_id})

@app.route('/api/tags/<int:tag_id>', methods=['PUT', 'DELETE'])
def handle_tag(tag_id):
    """Update or delete a specific tag."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    if request.method == 'PUT':
        data = request.json
        tag_name = data.get('tag_name')
        
        if not tag_name:
            return jsonify({'error': 'Tag name is required'}), 400
        
        cursor.execute('''
            UPDATE tags SET tag_name = %s WHERE tag_id = %s
        ''', (tag_name, tag_id))
        
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        # Check if tag is in use
        cursor.execute('SELECT COUNT(*) as count FROM property_tags WHERE tag_id = %s', (tag_id,))
        count = cursor.fetchone()['count']
        
        if count > 0:
            conn.close()
            return jsonify({'error': 'Cannot delete tag that is in use by properties'}), 400
        
        cursor.execute('DELETE FROM tags WHERE tag_id = %s', (tag_id,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})

# COMPANY INFORMATION ROUTES
@app.route('/company')
def company_management():
    """Company Information Page."""
    return render_template('company_management.html')

@app.route('/api/company', methods=['GET'])
def get_company():
    """Get company information (always returns record ID 1)."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    cursor.execute('SELECT * FROM companies WHERE c_id = 1')
    row = cursor.fetchone()
    
    if not row:
        # Create default company record if it doesn't exist
        cursor.execute('''
            INSERT INTO companies (c_id, c_name, c_phone, c_email)
            VALUES (1, 'Your Company Name', '(555) 123-4567', 'info@example.com')
            RETURNING c_id
        ''')
        conn.commit()
        cursor.execute('SELECT * FROM companies WHERE c_id = 1')
        row = cursor.fetchone()
    
    conn.close()
    return jsonify({'company': dict(row)})

@app.route('/api/company', methods=['PUT'])
def update_company():
    """Update company information."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    data = request.json
    
    # Update all fields including new c_url
    cursor.execute('''
        UPDATE companies SET
            c_phone = %s, c_fax = %s, c_email = %s, c_name = %s, c_url = %s,
            c_address = %s, c_city = %s, c_state = %s, c_zip = %s,
            c_nphone = %s, c_ophone = %s
        WHERE c_id = 1
    ''', (
        data.get('c_phone', ''), data.get('c_fax', ''), 
        data.get('c_email', ''), data.get('c_name', ''), data.get('c_url', ''),
        data.get('c_address', ''), data.get('c_city', ''), 
        data.get('c_state', ''), data.get('c_zip', ''),
        data.get('c_nphone', ''), data.get('c_ophone', '')
    ))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/company/signature', methods=['POST'])
def upload_signature():
    """Upload signature image file."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Validate file type
    allowed_types = {'image/png', 'image/jpeg', 'image/jpg'}
    if file.content_type not in allowed_types:
        return jsonify({'error': 'Only PNG or JPG files allowed'}), 400
    
    # Save file to uploads/company for sync capability
    filename = 'signature.png'
    filepath = os.path.join(STATIC_PATH, 'uploads', 'company', filename)
    
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    file.save(filepath)
    
    # Update database with file path
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('UPDATE companies SET c_sig_path = %s WHERE c_id = 1', 
                   (f'uploads/company/{filename}',))
    
    return jsonify({'success': True, 'filepath': f'photos/{filename}'})

# PROPERTY RECORD ROUTES
@app.route('/property/edit/<int:p_id>')
def edit_property(p_id):
    """Edit an existing property record."""
    return render_template('property_record.html', mode='edit', p_id=p_id)

@app.route('/api/properties/<int:p_id>')
def get_property(p_id):
    """Get full property data including related records."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get property with owner and status
    cursor.execute('''
        SELECT p.*, 
               s.s_status as p_status,
               CASE 
                   WHEN o.o_type = 'Company' THEN COALESCE(o.o_company, '')
                   ELSE TRIM(COALESCE(o.or_fname, '') || ' ' || COALESCE(o.or_lname, ''))
               END as or_name,
               o.or_fname, o.or_lname, o.or_email, o.or_phone, o.or_fax,
               o.o_type, o.o_fname, o.o_lname, o.o_2fname, o.o_2lname,
               o.o_3fname, o.o_3lname, o.o_4fname, o.o_4lname,
               o.o_5fname, o.o_5lname, o.o_company, o.or_m_address,
               o.or_m_address2, o.or_m_city, o.or_m_state, o.or_m_zip,
               o.o_other_owners
        FROM properties p
        JOIN owners o ON p.or_id = o.or_id
        LEFT JOIN statuses s ON p.p_status_id = s.status_id
        WHERE p.p_id = %s
    ''', (p_id,))
    
    property_data = dict(cursor.fetchone())
    
    # Get tags for this property
    cursor.execute('''
        SELECT t.tag_id, t.tag_name
        FROM property_tags pt
        JOIN tags t ON pt.tag_id = t.tag_id
        WHERE pt.p_id = %s
        ORDER BY t.tag_name
    ''', (p_id,))
    
    property_data['tags'] = [dict(row) for row in cursor.fetchall()]
    
    # Get photos
    cursor.execute('''
        SELECT photo_id, file_path, file_name, caption, is_primary
        FROM property_photos
        WHERE p_id = %s
        ORDER BY is_primary DESC, upload_date DESC
    ''', (p_id,))
    
    property_data['photos'] = [dict(row) for row in cursor.fetchall()]
    
    # Get documents
    cursor.execute('''
        SELECT doc_id, file_path, file_name, doc_type, description
        FROM property_documents
        WHERE p_id = %s
        ORDER BY upload_date DESC
    ''', (p_id,))
    
    property_data['documents'] = [dict(row) for row in cursor.fetchall()]
    
    # Get links
    cursor.execute('''
        SELECT link_id, url, description
        FROM property_links
        WHERE p_id = %s
        ORDER BY added_date DESC
    ''', (p_id,))
    
    property_data['links'] = [dict(row) for row in cursor.fetchall()]
    
    conn.close()
    return jsonify(property_data)

@app.route('/property/new')
def add_property():
    """Add a new property record."""
    copy_from = request.args.get('copy_from', type=int)
    
    # Get owner data if copying
    owner_data = {}
    if copy_from:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute('''
            SELECT o.* FROM owners o
            JOIN properties p ON o.or_id = p.or_id
            WHERE p.p_id = %s
        ''', (copy_from,))
        row = cursor.fetchone()
        if row:
            owner_data = dict(row)
        conn.close()

    print(f"DEBUG: Owner data for copy: {owner_data}")
    return render_template('property_record.html', mode='new', copy_from=copy_from, owner_data=owner_data)

@app.route('/api/properties', methods=['POST'])
def create_property():
    """Create a new property record."""
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        data = request.json
        
        print(f"=== CREATE PROPERTY DEBUG ===")
        print(f"p_status_id from request: {data.get('p_status_id')}")
        print(f"Type: {type(data.get('p_status_id'))}")
        print(f"Full data: {data}")
        
        # DEBUG: Print received status
        print(f"=== CREATE PROPERTY RECEIVED p_status_id: {data.get('p_status_id')} ===")
        print(f"Full data keys: {list(data.keys())}")
        
        # Required fields for all cases
        required_fields = ['p_apn', 'p_county', 'or_m_address', 'or_m_city', 'or_m_state']
        
        # Only require fname/lname for individuals
        o_type = data.get('o_type', 'Individual')
        if o_type != 'Company':
            required_fields.extend(['or_fname', 'or_lname'])
        
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # Check for existing owner
        if o_type == 'Company':
            # For companies, match on company name
            cursor.execute('''
                SELECT or_id FROM owners 
                WHERE o_company = %s AND o_type = 'Company'
            ''', (data.get('o_company'),))
        else:
            # For individuals, match on owner name
            cursor.execute('''
                SELECT or_id FROM owners 
                WHERE o_fname = %s AND o_lname = %s AND o_type = 'Individual'
            ''', (data.get('o_fname'), data.get('o_lname')))
        
        owner_row = cursor.fetchone()
        
        if owner_row:
            # Existing owner found - return special response for confirmation
            or_id = owner_row['or_id']
            
            # Get owner's existing properties count
            cursor.execute('SELECT COUNT(*) as count FROM properties WHERE or_id = %s', (or_id,))
            prop_count = cursor.fetchone()['count']
            
            return jsonify({
                'confirm': True,
                'message': f'Owner already exists with {prop_count} properties. Use existing owner?',
                'owner_id': or_id,
                'owner_name': data.get('o_company') if o_type == 'Company' else f"{data.get('o_fname')} {data.get('o_lname')}"
            }), 409  # Conflict status code
        
        # No existing owner - create new owner record
        cursor.execute('''
            INSERT INTO owners (o_type, or_fname, or_lname, or_email, or_phone, or_fax,
                              o_fname, o_lname, o_2fname, o_2lname, o_3fname, o_3lname,
                              o_4fname, o_4lname, o_5fname, o_5lname, o_company,
                              or_m_address, or_m_address2, or_m_city, or_m_state, or_m_zip)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING or_id
        ''', (
            data.get('o_type', 'Individual'), data.get('or_fname'), data.get('or_lname'),
            data.get('or_email'), data.get('or_phone'), data.get('or_fax'),
            data.get('o_fname'), data.get('o_lname'), data.get('o_2fname'),
            data.get('o_2lname'), data.get('o_3fname'), data.get('o_3lname'),
            data.get('o_4fname'), data.get('o_4lname'), data.get('o_5fname'),
            data.get('o_5lname'), data.get('o_company'), data.get('or_m_address'),
            data.get('or_m_address2'), data.get('or_m_city'), data.get('or_m_state'),
            data.get('or_m_zip')
        ))
        or_id = cursor.fetchone()['or_id']
        
        # Create property
        timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')
        
        cursor.execute('''
            INSERT INTO properties (or_id, p_status_id, p_apn, p_county, p_state, 
                          p_longstate, p_short_legal, p_address, p_city, p_zip,
                          p_acres, p_sqft, p_terrain, p_zoning, p_use, p_use_code,
                          p_use_description, p_impact_fee, p_environmental, p_restrictions,
                          p_waste_system_requirement, p_water_system_requirement,
                          p_survey, p_flood, p_flood_description, p_base_tax, p_hoa,
                          p_liens, p_back_tax, p_county_assessed_value, 
                          p_county_market_value, p_comp_market_value, p_improvements,
                          p_power, p_access, p_owned, p_aquired, p_est_value,
                          p_min_acceptable_offer, p_price, p_max_offer_amount,
                          p_m_date, p_offer_accept_date, p_contract_expires_date,
                          p_comments, p_note, p_plat_map_link, p_viable, p_betty_score, p_create_time, 
                          p_last_updated, p_status_last_updated, p_listed, p_agent_name, 
                          p_agent_phone, p_purchased_on, p_purchase_amount,
                          p_purchase_closing_costs, p_closing_company_name_purchase, 
                          p_sold_on, p_sold_amount, p_buyer, p_sold_closing_costs, 
                          p_profit, p_closing_company_name_sale,
                          p_last_sold_date, p_last_sold_amount, p_last_transaction_date, 
                          p_last_transaction_doc_type, p_mail_image_1, p_mail_image_2)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING p_id
        ''', (
            or_id, int(data.get('p_status_id')) if data.get('p_status_id') and str(data.get('p_status_id')).strip() != '' else None, data.get('p_apn'), data.get('p_county'),
            data.get('p_state'), get_longstate_from_abbr(data.get('p_state')), data.get('p_short_legal'),
            data.get('p_address'), data.get('p_city'), data.get('p_zip'),
            clean_numeric_field(data.get('p_acres')), 
            clean_numeric_field(data.get('p_sqft')), 
            data.get('p_terrain'), data.get('p_zoning'), data.get('p_use'), data.get('p_use_code'),
            data.get('p_use_description'), 
            clean_numeric_field(data.get('p_impact_fee')), 
            data.get('p_environmental'), data.get('p_restrictions'),
            data.get('p_waste_system_requirement'), data.get('p_water_system_requirement'),
            clean_boolean_field(data.get('p_survey')), data.get('p_flood'), data.get('p_flood_description'), 
            clean_numeric_field(data.get('p_base_tax')),
            clean_numeric_field(data.get('p_hoa')), 
            clean_numeric_field(data.get('p_liens')), 
            clean_numeric_field(data.get('p_back_tax')),
            clean_numeric_field(data.get('p_county_assessed_value')), 
            clean_numeric_field(data.get('p_county_market_value')),
            clean_numeric_field(data.get('p_comp_market_value')), data.get('p_improvements'),
            data.get('p_power'), data.get('p_access'), data.get('p_owned'),
            data.get('p_aquired'), 
            clean_numeric_field(data.get('p_est_value')), 
            clean_numeric_field(data.get('p_min_acceptable_offer')),
            clean_numeric_field(data.get('p_price')), 
            clean_numeric_field(data.get('p_max_offer_amount')), data.get('p_m_date'),
            data.get('p_offer_accept_date'), data.get('p_contract_expires_date'),
            data.get('p_comments'), data.get('p_note'), data.get('p_plat_map_link'), clean_boolean_field(data.get('p_viable')), clean_numeric_field(data.get('p_betty_score')),
            timestamp, timestamp, timestamp, clean_boolean_field(data.get('p_listed')), data.get('p_agent_name'),
            data.get('p_agent_phone'), data.get('p_purchased_on'), 
            clean_numeric_field(data.get('p_purchase_amount')),
            clean_numeric_field(data.get('p_purchase_closing_costs')),
            data.get('p_closing_company_name_purchase'), data.get('p_sold_on'), 
            clean_numeric_field(data.get('p_sold_amount')), data.get('p_buyer'), 
            clean_numeric_field(data.get('p_sold_closing_costs')),
            clean_numeric_field(data.get('p_profit')), data.get('p_closing_company_name_sale'),
            data.get('p_last_sold_date'), 
            clean_numeric_field(data.get('p_last_sold_amount')), 
            data.get('p_last_transaction_date'), 
            data.get('p_last_transaction_doc_type'),
            data.get('p_mail_image_1'), data.get('p_mail_image_2')
        ))
        
        p_id = cursor.fetchone()['p_id']
        
        # Insert tags
        for tag_id in data.get('tags', []):
            cursor.execute('INSERT INTO property_tags (p_id, tag_id) VALUES (%s, %s)', (p_id, tag_id))
        
        conn.commit()

        mark_record_modified('properties', p_id, conn)
        mark_record_modified('owners', or_id, conn)
        
        return jsonify({'success': True, 'p_id': p_id})
    except Exception as e:
        print(f"ERROR creating property: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/properties/confirm-create', methods=['POST'])
def confirm_create_property():
    """Create property using existing owner data but create NEW owner record if additional owners added"""
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        data = request.json
        
        # DEBUG: Print received status
        print(f"=== CONFIRM CREATE RECEIVED p_status_id: {data.get('p_status_id')} ===")

        # Get existing owner_id
        original_or_id = data.get('owner_id')
        
        if not original_or_id:
            return jsonify({'error': 'owner_id is required'}), 400
        
        # Check if user added new additional owners
        has_additional_owners = any([
            data.get('o_2fname'), data.get('o_2lname'),
            data.get('o_3fname'), data.get('o_3lname'),
            data.get('o_4fname'), data.get('o_4lname'),
            data.get('o_5fname'), data.get('o_5lname')
        ])
        
        if has_additional_owners:
            # NEW additional owners were added - create NEW owner record for isolation
            cursor.execute('''
                INSERT INTO owners (o_type, or_fname, or_lname, or_email, or_phone, or_fax,
                                  o_fname, o_lname, o_2fname, o_2lname, o_3fname, o_3lname,
                                  o_4fname, o_4lname, o_5fname, o_5lname, o_company,
                                  or_m_address, or_m_address2, or_m_city, or_m_state, or_m_zip)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING or_id
            ''', (
                data.get('o_type', 'Individual'), 
                data.get('or_fname'), data.get('or_lname'),
                data.get('or_email'), data.get('or_phone'), data.get('or_fax'),
                data.get('o_fname'), data.get('o_lname'), 
                data.get('o_2fname'), data.get('o_2lname'), 
                data.get('o_3fname'), data.get('o_3lname'),
                data.get('o_4fname'), data.get('o_4lname'), 
                data.get('o_5fname'), data.get('o_5lname'), 
                data.get('o_company'), data.get('or_m_address'),
                data.get('or_m_address2'), data.get('or_m_city'),
                data.get('or_m_state'), data.get('or_m_zip')
            ))
            or_id = cursor.fetchone()['or_id']
        else:
            # No additional owners - safe to reuse existing owner record
            or_id = original_or_id
        
        # Create property
        timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')
        
        cursor.execute('''
            INSERT INTO properties (or_id, p_status_id, p_apn, p_county, p_state, 
                          p_longstate, p_short_legal, p_address, p_city, p_zip,
                          p_acres, p_sqft, p_terrain, p_zoning, p_use, p_use_code,
                          p_use_description, p_impact_fee, p_environmental, p_restrictions,
                          p_waste_system_requirement, p_water_system_requirement,
                          p_survey, p_flood, p_flood_description, p_base_tax, p_hoa,
                          p_liens, p_back_tax, p_county_assessed_value, 
                          p_county_market_value, p_comp_market_value, p_improvements,
                          p_power, p_access, p_owned, p_aquired, p_est_value,
                          p_min_acceptable_offer, p_price, p_max_offer_amount,
                          p_m_date, p_offer_accept_date, p_contract_expires_date,
                          p_comments, p_note, p_plat_map_link, p_viable, p_betty_score, p_create_time, 
                          p_last_updated, p_status_last_updated, p_listed, p_agent_name, 
                          p_agent_phone, p_purchased_on, p_purchase_amount,
                          p_purchase_closing_costs, p_closing_company_name_purchase, 
                          p_sold_on, p_sold_amount, p_buyer, p_sold_closing_costs, 
                          p_profit, p_closing_company_name_sale,
                          p_last_sold_date, p_last_sold_amount, p_last_transaction_date, p_last_transaction_doc_type, 
                          p_mail_image_1, p_mail_image_2
                          )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING p_id
        ''', (
            or_id, int(data.get('p_status_id')) if data.get('p_status_id') and str(data.get('p_status_id')).strip() != '' else None, data.get('p_apn'), data.get('p_county'),
            data.get('p_state'), get_longstate_from_abbr(data.get('p_state')), data.get('p_short_legal'),
            data.get('p_address'), data.get('p_city'), data.get('p_zip'),
            clean_numeric_field(data.get('p_acres')), 
            clean_numeric_field(data.get('p_sqft')), 
            data.get('p_terrain'), data.get('p_zoning'), data.get('p_use'), data.get('p_use_code'),
            data.get('p_use_description'), 
            clean_numeric_field(data.get('p_impact_fee')), 
            data.get('p_environmental'), data.get('p_restrictions'),
            data.get('p_waste_system_requirement'), data.get('p_water_system_requirement'),
            clean_boolean_field(data.get('p_survey')), data.get('p_flood'), data.get('p_flood_description'), 
            clean_numeric_field(data.get('p_base_tax')),
            clean_numeric_field(data.get('p_hoa')), 
            clean_numeric_field(data.get('p_liens')), 
            clean_numeric_field(data.get('p_back_tax')),
            clean_numeric_field(data.get('p_county_assessed_value')), 
            clean_numeric_field(data.get('p_county_market_value')),
            clean_numeric_field(data.get('p_comp_market_value')), data.get('p_improvements'),
            data.get('p_power'), data.get('p_access'), data.get('p_owned'),
            data.get('p_aquired'), 
            clean_numeric_field(data.get('p_est_value')), 
            clean_numeric_field(data.get('p_min_acceptable_offer')),
            clean_numeric_field(data.get('p_price')), 
            clean_numeric_field(data.get('p_max_offer_amount')), data.get('p_m_date'),
            data.get('p_offer_accept_date'), data.get('p_contract_expires_date'),
            data.get('p_comments'), data.get('p_note'), data.get('p_plat_map_link'), clean_boolean_field(data.get('p_viable')), clean_numeric_field(data.get('p_betty_score')),
            timestamp, timestamp, timestamp, clean_boolean_field(data.get('p_listed')), data.get('p_agent_name'),
            data.get('p_agent_phone'), data.get('p_purchased_on'), 
            clean_numeric_field(data.get('p_purchase_amount')),
            clean_numeric_field(data.get('p_purchase_closing_costs')),
            data.get('p_closing_company_name_purchase'), data.get('p_sold_on'), 
            clean_numeric_field(data.get('p_sold_amount')), data.get('p_buyer'), 
            clean_numeric_field(data.get('p_sold_closing_costs')),
            clean_numeric_field(data.get('p_profit')), data.get('p_closing_company_name_sale'),
            data.get('p_last_sold_date'), 
            clean_numeric_field(data.get('p_last_sold_amount')), 
            data.get('p_last_transaction_date'), 
            data.get('p_last_transaction_doc_type'),
            data.get('p_mail_image_1'), data.get('p_mail_image_2')
        ))
        
        p_id = cursor.fetchone()['p_id']
        
        # Insert tags
        for tag_id in data.get('tags', []):
            cursor.execute('INSERT INTO property_tags (p_id, tag_id) VALUES (%s, %s)', (p_id, tag_id))
        
        conn.commit()

        mark_record_modified('properties', p_id, conn)
        mark_record_modified('owners', or_id, conn)
        
        return jsonify({'success': True, 'p_id': p_id})
    except Exception as e:
        print(f"ERROR in confirm-create: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/properties/<int:p_id>', methods=['PUT'])
def update_property(p_id):
    """Update an existing property record."""
    conn = None
    try: 
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        data = request.json
        
        # DEBUG: Print received status
        print(f"\n=== UPDATE PROPERTY RECEIVED p_status_id: {data.get('p_status_id')} ===")

        # === DEBUG PRINT #1: What data is coming from the form? ===
        print(f"\n=== DEBUG: Form Data ===")
        print(f"p_status_id from form: {data.get('p_status_id')} (type: {type(data.get('p_status_id'))})")
                
        # Build required fields based on owner type
        required_fields = ['p_apn', 'p_county', 'or_m_address', 'or_m_city', 'or_m_state']

        # Only require fname/lname for individuals
        o_type = data.get('o_type', 'Individual')
        if o_type != 'Company':
            required_fields.extend(['or_fname', 'or_lname'])

        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # FIX: Only update status timestamp if status actually changed
        # Get current status and timestamp from database
        cursor.execute('SELECT p_status_id, p_status_last_updated FROM properties WHERE p_id = %s', (p_id,))
        current_data = cursor.fetchone()
        
        current_status_id = current_data['p_status_id'] if current_data else None
        current_timestamp = current_data['p_status_last_updated'] if current_data else None
        new_status_id = data.get('p_status_id')
        
        # === DEBUG PRINT #2: What values are we comparing? ===
        print(f"Current DB status_id: {current_status_id} (type: {type(current_status_id)})")
        print(f"New form status_id: {new_status_id} (type: {type(new_status_id)})")
        
        # FIX: Force type comparison (string vs int issue)
        if str(current_status_id) == str(new_status_id):
            status_timestamp = current_timestamp  # Keep existing timestamp
            print("=== Status unchanged - keeping old timestamp ===")
        else:
            status_timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')  # New timestamp
            print("=== Status CHANGED - creating new timestamp ===")
        
        print(f"Final timestamp to save: {status_timestamp}")
        
        # Update property (rest of fields)
        timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')
        
        cursor.execute('''
            UPDATE properties SET
                p_status_id = %s, p_status_last_updated = %s, p_apn = %s, p_county = %s, p_state = %s, p_longstate = %s,
                p_short_legal = %s, p_address = %s, p_city = %s, p_zip = %s, p_acres = %s, 
                p_sqft = %s, p_terrain = %s, p_zoning = %s, p_use = %s, p_use_code = %s,
                p_use_description = %s, p_impact_fee = %s, p_environmental = %s, 
                p_restrictions = %s, p_waste_system_requirement = %s, 
                p_water_system_requirement = %s, p_survey = %s, p_flood = %s, p_flood_description = %s, p_base_tax = %s, p_hoa = %s, p_liens = %s, p_back_tax = %s, 
                p_county_assessed_value = %s, p_county_market_value = %s, 
                p_comp_market_value = %s, p_improvements = %s, p_power = %s, p_access = %s,
                p_owned = %s, p_aquired = %s, p_est_value = %s, p_min_acceptable_offer = %s,
                p_price = %s, p_max_offer_amount = %s, p_m_date = %s, 
                p_offer_accept_date = %s, p_contract_expires_date = %s, p_comments = %s,
                p_note = %s, p_plat_map_link = %s, p_viable = %s, p_betty_score = %s, p_last_updated = %s, p_listed = %s,
                p_agent_name = %s, p_agent_phone = %s, p_purchased_on = %s, 
                p_purchase_amount = %s, p_purchase_closing_costs = %s,
                p_closing_company_name_purchase = %s, p_sold_on = %s, p_sold_amount = %s,
                p_buyer = %s, p_sold_closing_costs = %s, p_profit = %s,
                p_closing_company_name_sale = %s,
                p_last_sold_date = %s, p_last_sold_amount = %s, 
                p_last_transaction_date = %s, p_last_transaction_doc_type = %s, p_mail_image_1 = %s, p_mail_image_2 = %s
            WHERE p_id = %s
        ''', (
            data.get('p_status_id'), status_timestamp, data.get('p_apn'), data.get('p_county'), 
            data.get('p_state'), get_longstate_from_abbr(data.get('p_state')), data.get('p_short_legal'),
            data.get('p_address'), data.get('p_city'), data.get('p_zip'), 
            clean_numeric_field(data.get('p_acres')), 
            clean_numeric_field(data.get('p_sqft')), data.get('p_terrain'),
            data.get('p_zoning'), data.get('p_use'), data.get('p_use_code'),
            data.get('p_use_description'), 
            clean_numeric_field(data.get('p_impact_fee')), 
            data.get('p_environmental'), data.get('p_restrictions'),
            data.get('p_waste_system_requirement'), data.get('p_water_system_requirement'),
            clean_boolean_field(data.get('p_survey')), data.get('p_flood'), data.get('p_flood_description'), 
            clean_numeric_field(data.get('p_base_tax')),
            clean_numeric_field(data.get('p_hoa')), 
            clean_numeric_field(data.get('p_liens')), 
            clean_numeric_field(data.get('p_back_tax')),
            clean_numeric_field(data.get('p_county_assessed_value')), 
            clean_numeric_field(data.get('p_county_market_value')),
            clean_numeric_field(data.get('p_comp_market_value')), data.get('p_improvements'), 
            data.get('p_power'), data.get('p_access'), data.get('p_owned'),
            data.get('p_aquired'), 
            clean_numeric_field(data.get('p_est_value')), 
            clean_numeric_field(data.get('p_min_acceptable_offer')),
            clean_numeric_field(data.get('p_price')), 
            clean_numeric_field(data.get('p_max_offer_amount')), data.get('p_m_date'),
            data.get('p_offer_accept_date'), data.get('p_contract_expires_date'),
            data.get('p_comments'), data.get('p_note'), data.get('p_plat_map_link'), clean_boolean_field(data.get('p_viable')), clean_numeric_field(data.get('p_betty_score')),
            timestamp, clean_boolean_field(data.get('p_listed')), data.get('p_agent_name'),
            data.get('p_agent_phone'), data.get('p_purchased_on'), 
            clean_numeric_field(data.get('p_purchase_amount')), 
            clean_numeric_field(data.get('p_purchase_closing_costs')),
            data.get('p_closing_company_name_purchase'), data.get('p_sold_on'),
            clean_numeric_field(data.get('p_sold_amount')), data.get('p_buyer'), 
            clean_numeric_field(data.get('p_sold_closing_costs')),
            clean_numeric_field(data.get('p_profit')), data.get('p_closing_company_name_sale'),
            data.get('p_last_sold_date'), 
            clean_numeric_field(data.get('p_last_sold_amount')), 
            data.get('p_last_transaction_date'), 
            data.get('p_last_transaction_doc_type'),
            data.get('p_mail_image_1'), data.get('p_mail_image_2'),
            p_id
        ))
        
        # Update owner record - with shared owner bug fix
        # Check if additional owners are being edited
        cursor.execute('''
            SELECT o_2fname, o_2lname, o_3fname, o_3lname,
                   o_4fname, o_4lname, o_5fname, o_5lname
            FROM owners WHERE or_id = (SELECT or_id FROM properties WHERE p_id = %s)
        ''', (p_id,))
        current_owner = cursor.fetchone()

        # Check if additional owners have changed
        additional_owners_changed = False
        if current_owner:
            new_values = [
                data.get('o_2fname'), data.get('o_2lname'),
                data.get('o_3fname'), data.get('o_3lname'),
                data.get('o_4fname'), data.get('o_4lname'),
                data.get('o_5fname'), data.get('o_5lname')
            ]
            current_values = [
                current_owner['o_2fname'], current_owner['o_2lname'],
                current_owner['o_3fname'], current_owner['o_3lname'],
                current_owner['o_4fname'], current_owner['o_4lname'],
                current_owner['o_5fname'], current_owner['o_5lname']
            ]
            
            # Compare values (handle None vs empty string)
            additional_owners_changed = any(
                (new or '') != (current or '') 
                for new, current in zip(new_values, current_values)
            )

        if additional_owners_changed:
            # Create new owner record to prevent affecting other properties
            cursor.execute('''
                INSERT INTO owners (
                    o_type, or_fname, or_lname, or_email, or_phone, or_fax,
                    o_fname, o_lname, o_2fname, o_2lname, o_3fname, o_3lname,
                    o_4fname, o_4lname, o_5fname, o_5lname, o_company,
                    or_m_address, or_m_address2, or_m_city, or_m_state, or_m_zip
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING or_id
            ''', (
                data.get('o_type'), data.get('or_fname'), data.get('or_lname'),
                data.get('or_email'), data.get('or_phone'), data.get('or_fax'),
                data.get('o_fname'), data.get('o_lname'), data.get('o_2fname'),
                data.get('o_2lname'), data.get('o_3fname'), data.get('o_3lname'),
                data.get('o_4fname'), data.get('o_4lname'), data.get('o_5fname'),
                data.get('o_5lname'), data.get('o_company'), data.get('or_m_address'),
                data.get('or_m_address2'), data.get('or_m_city'),
                data.get('or_m_state'), data.get('or_m_zip')
            ))
            new_or_id = cursor.fetchone()['or_id']
            
            # Update property to use new owner record
            cursor.execute('UPDATE properties SET or_id = %s WHERE p_id = %s', (new_or_id, p_id))
        else:
            # Update existing owner record (no additional owners changed)
            cursor.execute('''
                UPDATE owners SET
                    o_type = %s, or_fname = %s, or_lname = %s, or_email = %s, or_phone = %s, or_fax = %s,
                    o_fname = %s, o_lname = %s, o_2fname = %s, o_2lname = %s, o_3fname = %s,
                    o_3lname = %s, o_4fname = %s, o_4lname = %s, o_5fname = %s, o_5lname = %s,
                    o_company = %s, or_m_address = %s, or_m_address2 = %s, or_m_city = %s,
                    or_m_state = %s, or_m_zip = %s
                WHERE or_id = (SELECT or_id FROM properties WHERE p_id = %s)
            ''', (
                data.get('o_type'), data.get('or_fname'), data.get('or_lname'),
                data.get('or_email'), data.get('or_phone'), data.get('or_fax'),
                data.get('o_fname'), data.get('o_lname'), data.get('o_2fname'),
                data.get('o_2lname'), data.get('o_3fname'), data.get('o_3lname'),
                data.get('o_4fname'), data.get('o_4lname'), data.get('o_5fname'),
                data.get('o_5lname'), data.get('o_company'), data.get('or_m_address'),
                data.get('or_m_address2'), data.get('or_m_city'),
                data.get('or_m_state'), data.get('or_m_zip'), p_id
            ))
        
        # Update tags - delete all then re-insert
        cursor.execute('DELETE FROM property_tags WHERE p_id = %s', (p_id,))
        for tag_id in data.get('tags', []):
            cursor.execute('INSERT INTO property_tags (p_id, tag_id) VALUES (%s, %s)', (p_id, tag_id))
        
        conn.commit()

        # Mark for sync
        mark_record_modified('properties', p_id, conn)
        
        return jsonify({'success': True})

    except Exception as e:
        print(f"ERROR updating property: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

# UPLOAD ROUTES FOR PHOTOS/DOCUMENTS
@app.route('/api/properties/<int:p_id>/photos', methods=['POST'])
def upload_photo(p_id):
    """Upload a photo for a property."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = f"{p_id}_{file.filename}"
    relative_path = f'uploads/photos/p_{p_id}/{filename}'
    filepath = os.path.join(STATIC_PATH, relative_path)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    file.save(filepath)
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Insert property_photos record
        cursor.execute('''
            INSERT INTO property_photos (p_id, file_path, file_name, upload_date)
            VALUES (%s, %s, %s, NOW())
            RETURNING photo_id
        ''', (p_id, relative_path.replace('\\', '/'), filename))
        
        photo_id = cursor.fetchone()[0]
        
        # CRITICAL: Immediately add to file_sync for cloud upload
        try:
            from sync_service import ValiantLandSync
            sync = ValiantLandSync()
            file_hash = sync._get_file_hash(filepath)
            cursor.execute("""
                INSERT INTO file_sync (local_path, cloud_path, file_hash, modified_at, sync_status, last_sync_at)
                VALUES (%s, %s, %s, NOW(), 'pending', NOW())
                ON CONFLICT (local_path) DO UPDATE SET
                file_hash = EXCLUDED.file_hash,
                sync_status = 'pending',
                modified_at = NOW()
            """, (relative_path.replace('\\', '/'), None, file_hash))
        except Exception as e:
            print(f"Warning: Could not queue file for sync: {e}")
        
        # Mark property as modified FIRST (uses the transaction)
        mark_record_modified('properties', p_id, conn)
        
        # THEN commit everything together
        conn.commit()
        
        return jsonify({'success': True, 'photo_id': photo_id})
        
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/properties/<int:p_id>/documents', methods=['POST'])
def upload_document(p_id):
    """Upload a document for a property."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Generate filename with p_id
    filename = f"{p_id}_{file.filename}"
    relative_path = f'uploads/documents/p_{p_id}/{filename}'
    filepath = os.path.join(STATIC_PATH, relative_path)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    file.save(filepath)
    
    # Save to database
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Insert property_documents record
        cursor.execute('''
            INSERT INTO property_documents (p_id, file_path, file_name, upload_date)
            VALUES (%s, %s, %s, NOW())
            RETURNING doc_id
        ''', (p_id, relative_path.replace('\\', '/'), filename))
        
        doc_id = cursor.fetchone()[0]
        
        # CRITICAL: Immediately add to file_sync for cloud sync
        try:
            from sync_service import ValiantLandSync
            sync = ValiantLandSync()
            file_hash = sync._get_file_hash(filepath)
            cursor.execute("""
                INSERT INTO file_sync (local_path, cloud_path, file_hash, modified_at, sync_status, last_sync_at)
                VALUES (%s, %s, %s, NOW(), 'pending', NOW())
                ON CONFLICT (local_path) DO UPDATE SET
                file_hash = EXCLUDED.file_hash,
                sync_status = 'pending',
                modified_at = NOW()
            """, (relative_path.replace('\\', '/'), None, file_hash))
        except Exception as e:
            print(f"Warning: Could not queue document for sync: {e}")
        
        # FIX: Mark property modified BEFORE committing
        mark_record_modified('properties', p_id, conn)
        
        # THEN commit everything together
        conn.commit()
        
        return jsonify({'success': True, 'doc_id': doc_id})
        
    except Exception as e:
        conn.rollback()
        print(f"ERROR uploading document: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()
    
@app.route('/api/properties/<int:p_id>/links', methods=['POST'])
def add_link(p_id):
    """Add a link for a property."""
    data = request.json
    
    url = data.get('url')
    description = data.get('description', '')
    
    if not url:
        return jsonify({'error': 'URL is required'}), 400
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO property_links (p_id, url, description, added_date)
        VALUES (%s, %s, %s, NOW())
        RETURNING link_id
    ''', (p_id, url, description))
    
    link_id = cursor.fetchone()[0]
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'link_id': link_id})
    
# FIX 3: Modify upload_mail_image to track in file_sync (it already saves to disk, needs cloud sync tracking)
@app.route('/api/properties/<int:p_id>/mail-images', methods=['POST'])
def upload_mail_image(p_id):
    """Upload a mail image for a property (max 2)."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT p_mail_image_1, p_mail_image_2 FROM properties WHERE p_id = %s', (p_id,))
    result = cursor.fetchone()
    
    if not result:
        conn.close()
        return jsonify({'error': 'Property not found'}), 404
    
    if result[0] and result[1]:
        conn.close()
        return jsonify({'error': 'Maximum 2 mail images allowed. Please delete one first.'}), 400
    
    filename = f"{p_id}_{file.filename}"
    filepath = os.path.join(STATIC_PATH, 'uploads', 'photos', f'p_{p_id}', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    file.save(filepath)
    
    relative_path = f'uploads/photos/p_{p_id}/{filename}'
    
    if not result[0]:
        cursor.execute('UPDATE properties SET p_mail_image_1 = %s WHERE p_id = %s', (relative_path, p_id))
        slot = 1
    else:
        cursor.execute('UPDATE properties SET p_mail_image_2 = %s WHERE p_id = %s', (relative_path, p_id))
        slot = 2
    
    # FIX: Add to file_sync table so it gets uploaded to cloud
    try:
        from sync_service import ValiantLandSync
        sync = ValiantLandSync()
        file_hash = sync._get_file_hash(filepath)
        cursor.execute("""
            INSERT INTO file_sync (local_path, cloud_path, file_hash, modified_at, sync_status, last_sync_at)
            VALUES (%s, %s, %s, NOW(), 'pending', NOW())
            ON CONFLICT (local_path) DO UPDATE SET
            file_hash = EXCLUDED.file_hash,
            sync_status = 'pending',
            modified_at = NOW()
        """, (relative_path.replace('\\', '/'), None, file_hash))  # Normalize path
    except Exception as e:
        print(f"Warning: Could not queue mail image for sync: {e}")
    
    # Mark property as modified for sync
    cursor.execute("""
        UPDATE properties 
        SET modified_at = NOW(), sync_status = 'pending' 
        WHERE p_id = %s
    """, (p_id,))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'slot': slot, 'file_path': relative_path, 'file_name': filename})


@app.route('/api/properties/<int:p_id>/mail-images', methods=['DELETE'])
def delete_mail_image(p_id):
    """Delete a mail image from a specific slot."""
    data = request.json
    slot = data.get('slot')
    
    if slot not in [1, 2]:
        return jsonify({'error': 'Invalid slot'}), 400
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    column = f'p_mail_image_{slot}'
    cursor.execute(f'SELECT {column} FROM properties WHERE p_id = %s', (p_id,))
    result = cursor.fetchone()
    
    if result and result[0]:
        file_path = os.path.join(STATIC_PATH, result[0])
        if os.path.exists(file_path):
            os.remove(file_path)
        cursor.execute(f'UPDATE properties SET {column} = NULL WHERE p_id = %s', (p_id,))
        conn.commit()
    
    conn.close()
    return jsonify({'success': True})

@app.route('/static/<path:filename>', endpoint='static')
def serve_static(filename):
    """Serve static files - checks persistent folder first, then bundled"""
    try:
        # Security: Prevent directory traversal attacks (both slash types)
        if '..' in filename or filename.startswith('/') or filename.startswith('\\'):
            return "Invalid path", 400
        
        # Normalize filename: convert URL forward slashes to OS-specific separators
        filename = os.path.normpath(filename.replace('/', os.sep))
        
        # Prevent accessing files outside static folders (e.g., via symlinks or ..)
        if filename.startswith('..') or filename.startswith(os.sep):
            return "Invalid path", 400
            
        # First check persistent folder (uploads, user data) - BASE_DIR/static
        persistent_path = os.path.abspath(os.path.join(STATIC_PATH, filename))
        static_path_abs = os.path.abspath(STATIC_PATH)
        
        # Security check: ensure resolved path is actually within STATIC_PATH
        # (prevents tricks like static/../../../Windows/System32)
        if not persistent_path.startswith(static_path_abs + os.sep) and persistent_path != static_path_abs:
            print(f"[SECURITY BLOCK] Path {persistent_path} escapes {static_path_abs}")
            return "Access denied", 403
            
        if os.path.exists(persistent_path) and os.path.isfile(persistent_path):
            return send_file(persistent_path)
        
        # Fall back to bundled assets (CSS, JS, logos) - BUNDLE_DIR/static
        bundle_static_path = os.path.join(BUNDLE_DIR, 'static')
        bundle_path = os.path.abspath(os.path.join(bundle_static_path, filename))
        bundle_path_abs = os.path.abspath(bundle_static_path)
        
        # Security check for bundle path too
        if not bundle_path.startswith(bundle_path_abs + os.sep) and bundle_path != bundle_path_abs:
            print(f"[SECURITY BLOCK] Bundle path {bundle_path} escapes {bundle_path_abs}")
            return "Access denied", 403
            
        if os.path.exists(bundle_path) and os.path.isfile(bundle_path):
            return send_file(bundle_path)
        
        # Log missing files for debugging but return 404
        print(f"[404] File not found: {filename}")
        print(f"  Checked persistent: {persistent_path}")
        print(f"  Checked bundled: {bundle_path}")
        return "File not found", 404
        
    except Exception as e:
        import traceback
        print(f"[ERROR in serve_static] {str(e)}")
        print(f"  Filename requested: {filename}")
        print(f"  STATIC_PATH: {STATIC_PATH}")
        print(f"  BUNDLE_DIR: {BUNDLE_DIR}")
        print(traceback.format_exc())
        return f"Server error accessing file: {str(e)}", 500
        
# FIX 1: Add cascade delete to delete_photo endpoint
@app.route('/api/properties/<int:p_id>/photos/<int:photo_id>', methods=['DELETE'])
def delete_photo(p_id, photo_id):
    """Delete a photo and its file."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get file path before deleting record
    cursor.execute('SELECT file_path FROM property_photos WHERE photo_id = %s AND p_id = %s', (photo_id, p_id))
    row = cursor.fetchone()
    
    if row:
        file_path = row['file_path']
        full_path = os.path.join(STATIC_PATH, file_path)
        
        # FIX: Delete from file_sync first (cascade delete)
        cursor.execute('DELETE FROM file_sync WHERE local_path = %s', (file_path,))
        
        cursor.execute('DELETE FROM property_photos WHERE photo_id = %s AND p_id = %s', (photo_id, p_id))
        conn.commit()
        
        # Delete file from disk
        if os.path.exists(full_path):
            os.remove(full_path)
    
    conn.close()
    return jsonify({'success': True})

# FIX 2: Add cascade delete to delete_document endpoint  
@app.route('/api/properties/<int:p_id>/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(p_id, doc_id):
    """Delete a document and its file."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    cursor.execute('SELECT file_path FROM property_documents WHERE doc_id = %s AND p_id = %s', (doc_id, p_id))
    row = cursor.fetchone()
    
    if row:
        file_path = row['file_path']
        full_path = os.path.join(STATIC_PATH, file_path)
        
        # FIX: Delete from file_sync first (cascade delete)
        cursor.execute('DELETE FROM file_sync WHERE local_path = %s', (file_path,))
        
        cursor.execute('DELETE FROM property_documents WHERE doc_id = %s AND p_id = %s', (doc_id, p_id))
        conn.commit()
        
        if os.path.exists(full_path):
            os.remove(full_path)
    
    conn.close()
    return jsonify({'success': True})

@app.route('/api/properties/<int:p_id>/links/<int:link_id>', methods=['DELETE'])
def delete_link(p_id, link_id):
    """Delete a link and propagate to cloud."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Insert tombstone for sync propagation
        cursor.execute('''
            INSERT INTO sync_deletions (table_name, record_id, sync_status)
            VALUES (%s, %s, %s)
            ON CONFLICT (table_name, record_id) DO UPDATE SET
            deleted_at = NOW(), sync_status = 'pending', cloud_deleted = FALSE
        ''', ('property_links', link_id, 'pending'))
        
        # Try immediate cloud delete (optional)
        try:
            from sync_service import ValiantLandSync
            sync = ValiantLandSync()
            sync.supabase.table('property_links').delete().eq('link_id', link_id).execute()
            cursor.execute('UPDATE sync_deletions SET cloud_deleted = TRUE WHERE table_name = %s AND record_id = %s', ('property_links', link_id))
        except Exception as cloud_error:
            print(f"Cloud delete queued for link {link_id}: {cloud_error}")
        
        # Delete local link
        cursor.execute('DELETE FROM property_links WHERE link_id = %s AND p_id = %s', (link_id, p_id))
        conn.commit()
        
        return jsonify({'success': True, 'message': 'Link deleted'})
        
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

# ============================================
# DOCUMENT GENERATION ROUTES
# ============================================

@app.route('/documents')
def document_creation():
    """Document creation page."""
    return render_template('document_creation.html')

@app.route('/api/documents/templates', methods=['GET'])
def get_templates():
    """Get all available document templates."""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    cursor.execute('SELECT template_id, name, filename FROM document_templates WHERE is_active = TRUE ORDER BY name')
    templates = [dict(row) for row in cursor.fetchall()]
    
    conn.close()
    return jsonify({'templates': templates})

@app.route('/api/documents/generate', methods=['POST'])
def generate_documents():
    """Generate documents for selected properties and return as direct download."""
    data = request.json
    template_id = data.get('template_id')
    property_ids = data.get('property_ids', [])
    
    if not template_id or not property_ids:
        return jsonify({'error': 'Template ID and property IDs are required'}), 400
    
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get template
        cursor.execute('SELECT * FROM document_templates WHERE template_id = %s', (template_id,))
        template = cursor.fetchone()
        
        if not template:
            return jsonify({'error': 'Template not found'}), 404
        
        # Get properties with owner data
        cursor.execute('''
            SELECT p.*, o.*, s.s_status as p_status
            FROM properties p
            JOIN owners o ON p.or_id = o.or_id
            LEFT JOIN statuses s ON p.p_status_id = s.status_id
            WHERE p.p_id IN %s
        ''', (tuple(property_ids),))
        
        properties = [dict(row) for row in cursor.fetchall()]
        
        # Get company data
        cursor.execute('SELECT * FROM companies WHERE c_id = 1')
        company = dict(cursor.fetchone())
        
        # Check if this is a Contract Template
        is_contract = template['name'] == 'Contract Template'
        
        # Generate filename prefix
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        prefix_map = {
            'Neutral Letter': 'Neutral_Letters',
            'Blind Offer Letter': 'BO_Letters',
            'Offer Letter': 'Offer_Letters',
            'Multi Offer Letter': 'Multi_Offer_Letters',
            '2nd Offer Letter': '2nd_Offer_Letters',
            'Multi 2nd Offer Letter': 'Multi_2nd_Offer_Letters',
            'Contract Template': 'Contracts',
            'Postcard Template': 'Postcard'
        }
        prefix = prefix_map.get(template['name'], 'Document')
        final_filename = f"{prefix}_{timestamp}.docx"
        
        master_doc = None
        
        if is_contract:
            # CONTRACT LOGIC - One document per property
            properties_sorted = sorted(properties, key=lambda x: (x['or_id'], x['p_id']))
            
            for i, prop in enumerate(properties_sorted):
                owner = {k: v for k, v in prop.items() if k.startswith('o_') or k.startswith('or_')}
                
                doc_content = generate_document_for_owner_fix(
                    owner, 
                    [prop],
                    template['filename'],
                    company
                )
                
                if i == 0:
                    master_doc = doc_content
                else:
                    append_document_correctly(master_doc, doc_content, is_contract=True, is_postcard=False)
                    
        else:
            # LETTER LOGIC - Group by owner
            owners_dict = {}
            for prop in properties:
                or_id = prop['or_id']
                if or_id not in owners_dict:
                    owners_dict[or_id] = {
                        'owner': {k: v for k, v in prop.items() if k.startswith('o_') or k.startswith('or_')},
                        'properties': []
                    }
                owners_dict[or_id]['properties'].append(
                    {k: v for k, v in prop.items() if not k.startswith('o_') and not k.startswith('or_')}
                )
            
            for i, (or_id, owner_data) in enumerate(owners_dict.items()):
                owner = owner_data['owner']
                props = owner_data['properties']
                
                template_to_use = select_template_by_property_count(template['name'], len(props), cursor)
                
                doc_content = generate_document_for_owner_fix(owner, props, template_to_use, company)
                
                if i == 0:
                    master_doc = doc_content
                else:
                    is_postcard = 'Postcard' in template_to_use
                    append_document_correctly(master_doc, doc_content, is_contract=False, is_postcard=is_postcard)
        
        # Save to BytesIO
        if master_doc:
            from io import BytesIO
            doc_buffer = BytesIO()
            master_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return send_file(
                doc_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=final_filename
            )
        else:
            return jsonify({'error': 'No documents generated'}), 400
        
    except Exception as e:
        print(f"ERROR generating documents: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

def select_template_by_property_count(base_template_name, prop_count, cursor):
    """Select appropriate template based on property count."""
    if prop_count == 1:
        # Query database for the single property template filename
        cursor.execute('SELECT filename FROM document_templates WHERE name = %s', (base_template_name,))
        result = cursor.fetchone()
        if result:
            return result['filename']
        else:
            return base_template_name + '.docx'
    
    # Multi-property templates
    template_map = {
        'Offer Letter': 'Multi Offer Letter',
        '2nd Offer Letter': 'Multi 2nd Offer Letter'
    }
    
    base_name = base_template_name.replace('.docx', '')
    multi_name = template_map.get(base_name, base_name)
    
    cursor.execute('SELECT filename FROM document_templates WHERE name = %s', (multi_name,))
    result = cursor.fetchone()
    
    if result:
        return result['filename']
    else:
        return base_template_name

def set_document_auto_update(doc):
    """Set document to auto-update fields when opened (prevents need for Ctrl+A, F9)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    settings = doc.settings._element
    # Create updateFields element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

def replace_bookmarked_image(doc, bookmark_name, image_path):
    """Replace image at bookmark with new image, preserving all formatting including wrapping"""
    from docx.oxml.ns import qn
    
    if not image_path or not os.path.exists(image_path):
        print(f"Image not found for bookmark {bookmark_name}: {image_path}")
        return False
    
    # Find all bookmark start elements
    body = doc._element.body
    
    # Method 1: Look for bookmarkStart directly
    bookmark_start = None
    for elem in body.iter():
        if elem.tag == qn('w:bookmarkStart'):
            name_attr = elem.get(qn('w:name'))
            if name_attr == bookmark_name:
                bookmark_start = elem
                break
    
    if bookmark_start is None:
        print(f"Bookmark {bookmark_name} start not found")
        return False
    
    # Find the bookmark ID
    bookmark_id = bookmark_start.get(qn('w:id'))
    
    # Now find the bookmarkEnd with the same ID
    bookmark_end = None
    for elem in body.iter():
        if elem.tag == qn('w:bookmarkEnd'):
            if elem.get(qn('w:id')) == bookmark_id:
                bookmark_end = elem
                break
    
    if bookmark_end is None:
        print(f"Bookmark {bookmark_name} end not found")
        return False
    
    # IMPROVED: Find the blip (image reference) within the bookmark range
    blip = None
    
    # Get parent of bookmark start
    start_parent = bookmark_start.getparent()
    
    # Strategy 1: If bookmark is within a single paragraph, search that paragraph
    if start_parent is not None:
        for elem in start_parent.iter():
            if elem == bookmark_start:
                continue
            if elem == bookmark_end:
                break
            if elem.tag == qn('a:blip'):
                blip = elem
                break
    
    # Strategy 2: If not found, search all paragraphs for one containing this bookmark
    if blip is None:
        for para in doc.paragraphs:
            para_elem = para._p
            has_bookmark = False
            
            for child in para_elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == bookmark_name:
                        has_bookmark = True
                        break
            
            if has_bookmark:
                # Search for blip in this paragraph
                for elem in para_elem.iter():
                    if elem.tag == qn('a:blip'):
                        blip = elem
                        break
                if blip:
                    break
    
    # Strategy 3: Search entire body between bookmarkStart and bookmarkEnd elements
    if blip is None:
        found_start = False
        for elem in body.iter():
            if elem == bookmark_start:
                found_start = True
                continue
            if found_start:
                if elem == bookmark_end:
                    break
                if elem.tag == qn('a:blip'):
                    blip = elem
                    break
    
    if blip is None:
        print(f"No blip found for bookmark {bookmark_name}")
        return False
    
    # Get the relationship ID
    embed = blip.get(qn('r:embed'))
    if not embed:
        print(f"No embed attribute found for bookmark {bookmark_name}")
        return False
    
    # Get the image part
    try:
        image_part = doc.part.related_parts[embed]
    except KeyError:
        print(f"Image part {embed} not found for bookmark {bookmark_name}")
        return False
    
    # Replace the image blob
    try:
        with open(image_path, 'rb') as f:
            new_blob = f.read()
        
        image_part._blob = new_blob
        print(f"Successfully replaced image at bookmark {bookmark_name}")
        return True
        
    except Exception as e:
        print(f"Error replacing image at bookmark {bookmark_name}: {e}")
        return False

def generate_document_for_owner_fix(owner, properties, template_filename, company):
    """Generate document for a single owner with proper field calculation and preservation."""
    template_path = get_template_path(template_filename)
    
    if not os.path.exists(template_path):
        raise Exception(f"Template not found: {template_filename}")
    
    # Load template
    doc = Document(template_path)
       
    # Prepare merge data
    merge_data = prepare_merge_data(owner, properties, company)
    
    # Replace bookmarked images first (before text replacement)
    # This preserves the "behind text" wrapping from the template
    # Only for postcard templates
    if 'Postcard' in template_filename:
        if merge_data.get('p_mail_image_1'):
            replace_bookmarked_image(doc, 'p_mail_image_1', merge_data['p_mail_image_1'])
        
        if merge_data.get('p_mail_image_2'):
            replace_bookmarked_image(doc, 'p_mail_image_2', merge_data['p_mail_image_2'])
    
    # Replace merge fields (text only - images are already handled)
    replace_merge_fields_fix(doc, merge_data)

    # Remove headers/footers for postcard templates to prevent layout issues
    if 'Postcard' in template_filename:
        print(f"DEBUG: Removing headers/footers for {template_filename}")
        for section in doc.sections:
            # Clear header content by removing all paragraphs
            header = section.header
            if header and len(header.paragraphs) > 0:
                # Remove each paragraph from the header
                for para in list(header.paragraphs):
                    p_element = para._p
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                print(f"DEBUG: Cleared {len(header.paragraphs)} paragraphs from header")
            
            # Clear footer content by removing all paragraphs  
            footer = section.footer
            if footer and len(footer.paragraphs) > 0:
                for para in list(footer.paragraphs):
                    p_element = para._p
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                print(f"DEBUG: Cleared {len(footer.paragraphs)} paragraphs from footer")
            
            # Also unlink to be safe
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
            
    else:
        pass
        
    return doc

def prepare_merge_data(owner, properties, company):
    """Prepare all merge fields with proper calculations, formatting, and empty line removal."""
    merge_data = {}
    
    # Add company data with c_ prefix
    for key, value in company.items():
        if value is None:
            value = ''
        clean_key = key[2:] if key.startswith('c_') else key
        merge_data[f'c_{clean_key}'] = str(value)
    
    # Add owner data (ensure no missing keys)
    for key, value in owner.items():
        if value is None:
            value = ''
        merge_data[key] = str(value)
    
    # Calculate or_name and or_greeting if not present
    o_type = owner.get('o_type', 'Individual')
    if 'or_name' not in merge_data or not merge_data['or_name']:
        if o_type == 'Company':
            merge_data['or_name'] = str(owner.get('o_company', ''))
        else:
            fname = str(owner.get('or_fname', ''))
            lname = str(owner.get('or_lname', ''))
            merge_data['or_name'] = f"{fname} {lname}".strip()
    
    if 'or_greeting' not in merge_data or not merge_data['or_greeting']:
        if o_type == 'Company':
            merge_data['or_greeting'] = "To whom it may concern,"
        else:
            fname = str(owner.get('or_fname', ''))
            merge_data['or_greeting'] = f"Dear {fname},"
    
    # CROSS-PLATFORM DATE FORMATTING FIX
    current_time = datetime.now()
    
    def format_date_no_leading_zero(date_obj):
        """Works on ALL platforms: Windows, macOS, Linux"""
        month = date_obj.strftime('%B')
        day = str(date_obj.day)
        year = date_obj.strftime('%Y')
        return f"{month} {day}, {year}"
    
    # Calculate property-specific fields
    is_multi = len(properties) > 1
    
    # Define calculated fields that should not be overwritten by database values
    calculated_fields = {'m_date', 'p_offer_accept_date', 'p_contract_expires_date', 'm_net', 
                        'p_price', 'p_back_tax', 'mo_p_price', 'mo_p_back_tax', 'mo_net'}
    
    if is_multi:
        # Multi-property calculations
        mo_p_price = sum(float(p.get('p_price', 0) or 0) for p in properties)
        mo_p_back_tax = sum(float(p.get('p_back_tax', 0) or 0) for p in properties)
        mo_net = mo_p_price - mo_p_back_tax
        
        # Use first property as base, but DON'T overwrite calculated fields
        prop = properties[0]
        for key, value in prop.items():
            if not key.startswith('c_') and not key.startswith('o_') and not key.startswith('or_'):
                if value is None:
                    value = ''
                if key not in calculated_fields:
                    merge_data[key] = str(value)
        
        # Set calculated multi-property values
        merge_data['mo_p_price'] = f"{mo_p_price:,.2f}"
        merge_data['mo_p_back_tax'] = f"{mo_p_back_tax:,.2f}"
        merge_data['mo_net'] = f"{mo_net:,.2f}"
        
        # Handle legacy template expression
        merge_data['mo_p_price-mo_p_backtax'] = f"{mo_net:,.2f}"
    else:
        # Single property calculations
        prop = properties[0]
        p_price = float(prop.get('p_price', 0) or 0)
        p_back_tax = float(prop.get('p_back_tax', 0) or 0)
        m_net = p_price - p_back_tax
        
        # Add property fields, but DON'T overwrite calculated fields
        for key, value in prop.items():
            if not key.startswith('c_') and not key.startswith('o_') and not key.startswith('or_'):
                if value is None:
                    value = ''
                if key not in calculated_fields:
                    merge_data[key] = str(value)
        
        # Set calculated single-property values
        merge_data['m_net'] = f"{m_net:,.2f}"
        merge_data['p_price'] = f"{p_price:,.2f}"
        merge_data['p_back_tax'] = f"{p_back_tax:,.2f}"
    
    # Set dates LAST to ensure they aren't overwritten
    merge_data['m_date'] = format_date_no_leading_zero(current_time)
    merge_data['p_offer_accept_date'] = format_date_no_leading_zero(current_time + timedelta(days=15))
    merge_data['p_contract_expires_date'] = format_date_no_leading_zero(current_time + timedelta(days=180))
    
    # Track empty fields that should be removed from document
    merge_data['_empty_fields'] = []
    for key, value in merge_data.items():
        if 'address2' in key and value.strip() == '':
            merge_data['_empty_fields'].append(key)
    
    # Convert mail image relative paths to absolute paths for Word INCLUDEPICTURE fields
    if merge_data.get('p_mail_image_1'):
        # Use forward slashes for Word compatibility and wrap in quotes for spaces
        abs_path = os.path.join(STATIC_PATH, merge_data['p_mail_image_1']).replace('\\', '/')
        merge_data['p_mail_image_1'] = abs_path
    if merge_data.get('p_mail_image_2'):
        abs_path = os.path.join(STATIC_PATH, merge_data['p_mail_image_2']).replace('\\', '/')
        merge_data['p_mail_image_2'] = abs_path 
    
    return merge_data

def get_all_paragraphs(doc):
    """Get all paragraphs including those in text boxes, tables, headers, and footers"""
    from docx.text.paragraph import Paragraph
    
    paragraphs = []
    seen = set()
    
    def add_paragraph(p_elem, part):
        if id(p_elem) not in seen:
            seen.add(id(p_elem))
            try:
                paragraphs.append(Paragraph(p_elem, part))
            except:
                pass
    
    # Standard paragraphs
    for p in doc.paragraphs:
        add_paragraph(p._element, doc._body)
    
    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    add_paragraph(p._element, doc._body)
    
    # Headers and footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                add_paragraph(p._element, section.header.part)
        if section.footer:
            for p in section.footer.paragraphs:
                add_paragraph(p._element, section.footer.part)
    
    # TEXT BOXES: Find all w:p elements anywhere in document XML
    # Text boxes contain paragraphs not in the main flow
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Search entire document body (catches text boxes anchored to body)
    for p_elem in doc._element.body.findall('.//w:p', nsmap):
        add_paragraph(p_elem, doc._body)
    
    # Search headers/footers for text boxes
    for section in doc.sections:
        if section.header:
            for p_elem in section.header._element.findall('.//w:p', nsmap):
                add_paragraph(p_elem, section.header.part)
        if section.footer:
            for p_elem in section.footer._element.findall('.//w:p', nsmap):
                add_paragraph(p_elem, section.footer.part)
    
    return paragraphs

def replace_merge_fields_fix(doc, merge_data):
    """Replace placeholders while preserving paragraph formatting and removing empty lines."""
    
    def process_paragraph(paragraph):
        """Process a single paragraph without destroying formatting."""
        if '[' not in paragraph.text:
            return
        
        text = paragraph.text
        
        # Skip image fields - they are handled by bookmarks now
        if '[p_mail_image_1]' in text or '[p_mail_image_2]' in text:
            # Remove the placeholder text but keep paragraph for layout
            # Or remove paragraph entirely if it only contains the placeholder
            cleaned = text.replace('[p_mail_image_1]', '').replace('[p_mail_image_2]', '').strip()
            if not cleaned:
                # Remove empty paragraph
                p_element = paragraph._p
                parent = p_element.getparent()
                if parent is not None:
                    parent.remove(p_element)
            else:
                # Replace with cleaned text
                for run in list(paragraph.runs):
                    if run._r.getparent() is not None:
                        run._r.getparent().remove(run._r)
                paragraph.add_run(cleaned)
            return
        
        # Regular text replacement for non-image fields
        full_text = text
        for field_name, value in merge_data.items():
            if field_name.startswith('p_mail_image'):
                continue  # Skip image fields
            placeholder = f'[{field_name}]'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # Enhanced empty detection
        cleaned_text = full_text.strip()
        
        if cleaned_text == '':
            p_element = paragraph._p
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            return
        
        # Clear runs properly
        for run in list(paragraph.runs):
            if run._r.getparent() is not None:
                run._r.getparent().remove(run._r)
        
        if full_text:
            paragraph.add_run(full_text)
        elif full_text.strip() == '':
            p_element = paragraph._p
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
    
    # Process all paragraphs
    for paragraph in list(get_all_paragraphs(doc)):
        process_paragraph(paragraph)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    process_paragraph(paragraph)
    
    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in list(section.header.paragraphs):
                process_paragraph(paragraph)
        if section.footer:
            for paragraph in list(section.footer.paragraphs):
                process_paragraph(paragraph)

def convert_image_placeholders_to_fields(doc):
    """Convert image path text to Word INCLUDEPICTURE fields"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Process ALL paragraphs including text boxes
    for paragraph in get_all_paragraphs(doc):
        text = paragraph.text
        
        # Skip if no potential path
        if not text or 'static' not in text:
            continue
            
        # Normalize text - remove newlines that split paths across lines
        full_text = text.replace('\n', ' ').replace('\r', '').strip()
        
        # Find ALL image paths in this paragraph (non-greedy match, allows spaces)
        # Matches: C:/.../static/uploads/photos/p_12345/filename.jpg (or .png, etc.)
        matches = re.finditer(
            r'([A-Z]:[\\\/].*?static[\\\/]uploads[\\\/]photos[\\\/].*?\.(?:jpg|jpeg|png|gif|bmp))',
            full_text,
            re.IGNORECASE
        )
        
        for match in matches:
            image_path = match.group(1)
            
            # Clear paragraph and create INCLUDEPICTURE field
            paragraph.clear()
            run = paragraph.add_run()
            
            # Build field structure: { INCLUDEPICTURE "path" \\* MERGEFORMAT \\d }
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = f' INCLUDEPICTURE "{image_path}" \\\\* MERGEFORMAT \\\\d '
            
            fldChar_separate = OxmlElement('w:fldChar')
            fldChar_separate.set(qn('w:fldCharType'), 'separate')
            
            # Add some placeholder text (shows if image not found)
            placeholder = OxmlElement('w:t')
            placeholder.text = "[Image]"
            
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar_begin)
            run._r.append(instrText)
            run._r.append(fldChar_separate)
            run._r.append(placeholder)  # Optional: shows if image missing
            run._r.append(fldChar_end)
            
            # Only process first match per paragraph for safety
            # If you have multiple images per paragraph, this needs refactoring
            break

def append_document_correctly(target_doc, source_doc, is_contract=False, is_postcard=False):
    """PROPERLY append source_doc to target_doc preserving images and formatting."""
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    import copy
    
    # Save source to temp stream
    source_stream = BytesIO()
    source_doc.save(source_stream)
    source_stream.seek(0)
    
    # Load fresh copy
    temp_doc = Document(source_stream)
    
    # Copy images
    _copy_document_images(temp_doc, target_doc)
    
    def element_has_content(element):
        """Check if element has text or images/drawings"""
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        }
        has_text = any(t.text and t.text.strip() for t in element.iter() if hasattr(t, 'text'))
        has_image = (
            element.find('.//w:drawing', nsmap) is not None or 
            element.find('.//w:pict', nsmap) is not None or
            element.find('.//w:object', nsmap) is not None
        )
        return has_text or has_image
    
    # Get body elements
    body_elements = list(temp_doc._element.body)
    
    # Find start index (skip leading empty paragraphs)
    start_idx = 0
    for i, element in enumerate(body_elements):
        if element.tag.endswith('p'):
            if element_has_content(element):
                start_idx = i
                break
        else:
            start_idx = i
            break
    
    # Find end index (skip trailing sectPr and empty paragraphs)
    end_idx = len(body_elements)
    
    # Skip final sectPr
    if end_idx > 0 and body_elements[end_idx-1].tag.endswith('sectPr'):
        end_idx -= 1
    
    # Skip trailing empty paragraphs (CRITICAL: prevents blank pages)
    for i in range(end_idx-1, start_idx-1, -1):
        element = body_elements[i]
        if element.tag.endswith('p'):
            if element_has_content(element):
                break
            else:
                end_idx -= 1
        else:
            break
    
    # Only add page break if target already has content
    has_content = any(p.text.strip() for p in target_doc.paragraphs) or len(target_doc._element.body) > 1
    if has_content:
        last_para = target_doc.paragraphs[-1] if target_doc.paragraphs else target_doc.add_paragraph()
        last_para.add_run().add_break(WD_BREAK.PAGE)
    
    # Append only the content elements (no trailing garbage)
    for element in body_elements[start_idx:end_idx]:
        copied = copy.deepcopy(element)
        target_doc._element.body.append(copied)
    
    # Remove headers/footers for contracts and postcards only (letters keep headers)
    if is_contract or is_postcard:
        for section in target_doc.sections:
            try:
                section.header.is_linked_to_previous = True
                section.footer.is_linked_to_previous = True
            except:
                pass
    
    source_stream.close()

def _copy_document_images(source_doc, target_doc):
    """Copy image binaries from source to target document"""
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    import hashlib
    
    try:
        source_part = source_doc.part
        target_part = target_doc.part
        
        # Find all blips (image references) in source document
        # Search in main document part, headers, and footers
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        # Get all blips from the source part's XML
        source_blips = source_part.element.findall('.//a:blip', namespaces)
        
        for blip in source_blips:
            embed = blip.get(qn('r:embed'))
            if not embed:
                continue
                
            try:
                # Get the image part from source
                source_image_part = source_part.related_parts.get(embed)
                if not source_image_part:
                    continue
                
                # Generate unique partname based on content hash
                content_hash = hashlib.md5(source_image_part.blob).hexdigest()[:8]
                ext = source_image_part.partname.split('.')[-1].lower()
                if ext not in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                    ext = 'png'
                new_partname_str = f'/word/media/image_{content_hash}.{ext}'
                
                # Check if this image already exists in target by comparing string representations
                existing_part = None
                for part in target_part.package.parts:
                    if str(part.partname) == new_partname_str:
                        existing_part = part
                        break
                
                if existing_part:
                    new_image_part = existing_part
                else:
                    # Determine content type
                    content_types = {
                        'png': 'image/png', 'jpg': 'image/jpeg', 
                        'jpeg': 'image/jpeg', 'gif': 'image/gif', 'bmp': 'image/bmp'
                    }
                    ct = content_types.get(ext, 'image/png')
                    
                    # Create PackURI object (not a string!)
                    new_partname = PackURI(new_partname_str)
                    
                    # Create new part with proper PackURI
                    from docx.opc.part import Part
                    new_image_part = Part(new_partname, ct, source_image_part.blob, target_part.package)
                    
                    # Add to package parts list
                    # Access the internal _parts list to append
                    if hasattr(target_part.package, '_parts'):
                        target_part.package._parts.append(new_image_part)
                    else:
                        # Fallback: try to access parts collection
                        parts_collection = target_part.package.parts
                        if hasattr(parts_collection, '_parts'):
                            parts_collection._parts.append(new_image_part)
                        elif hasattr(parts_collection, 'append'):
                            parts_collection.append(new_image_part)
                
                # Generate unique rId in target
                existing_rids = set(target_part.rels.keys())
                counter = 1
                while f'rId{counter}' in existing_rids:
                    counter += 1
                new_rId = f'rId{counter}'
                
                # Add relationship from target part to image part
                target_part.rels.add_relationship(RT.IMAGE, new_image_part, new_rId)
                
                # Update the blip to use the new rId
                # This modifies the source document's element, which will be copied to target
                blip.set(qn('r:embed'), new_rId)
                        
            except Exception as e:
                print(f"Image copy warning: {e}")
                import traceback
                traceback.print_exc()
                continue
                
    except Exception as e:
        print(f"Image processing warning: {e}")
        import traceback
        traceback.print_exc()

def _copy_paragraph_content(target_para, source_para):
    """Copy paragraph content including runs with proper formatting"""
    from copy import deepcopy
    
    # Copy paragraph formatting
    target_para._p.get_or_add_pPr().set(source_para._p.get_or_add_pPr().xml)
    
    # Copy each run
    for run in source_para.runs:
        new_run = target_para.add_run()
        new_run.text = run.text
        if run._element.rPr is not None:
            new_run._element.get_or_add_rPr().set(run._element.rPr.xml)

@app.route('/api/documents/progress/<progress_id>', methods=['GET'])
def get_progress(progress_id):
    """Get generation progress."""
    progress = generation_progress.get(progress_id, {'status': 'not_found'})
    return jsonify(progress)
    
@app.route('/api/properties/mark-as-mailed', methods=['POST'])
def mark_as_mailed():
    """Mark properties as mailed by setting p_m_date, p_offer_accept_date, and p_contract_expires_date."""
    conn = None
    try:
        data = request.json
        property_ids = data.get('property_ids', [])
        
        if not property_ids:
            return jsonify({'error': 'No property IDs provided'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Calculate dates based on click timestamp
        today = datetime.now()
        p_m_date = today.strftime('%Y-%m-%d')
        p_offer_accept_date = (today + timedelta(days=15)).strftime('%Y-%m-%d')
        p_contract_expires_date = (today + timedelta(days=180)).strftime('%Y-%m-%d')
        
        # Update all selected properties
        placeholders = ','.join(['%s'] * len(property_ids))
        cursor.execute(f'''
            UPDATE properties 
            SET p_m_date = %s,
                p_offer_accept_date = %s,
                p_contract_expires_date = %s,
                p_last_updated = %s
            WHERE p_id IN ({placeholders})
        ''', [p_m_date, p_offer_accept_date, p_contract_expires_date, 
              today.strftime('%m/%d/%Y %I:%M %p')] + property_ids)
        
        conn.commit()

        for p_id in property_ids:
            mark_record_modified('properties', p_id, conn)
        
        return jsonify({
            'success': True,
            'updated': cursor.rowcount,
            'dates': {
                'p_m_date': p_m_date,
                'p_offer_accept_date': p_offer_accept_date,
                'p_contract_expires_date': p_contract_expires_date
            }
        })
        
    except Exception as e:
        print(f"ERROR marking as mailed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/documents/confirm', methods=['POST'])
def confirm_documents():
    """Confirm and save generated documents."""
    data = request.json
    progress_id = data.get('progress_id')
    
    if not progress_id:
        return jsonify({'error': 'Progress ID required'}), 400
    
    progress = generation_progress.get(progress_id)
    if not progress or progress['status'] != 'complete':
        return jsonify({'error': 'Invalid or incomplete generation'}), 400
    
    # Move from temp to permanent location
    temp_filename = progress['filename']
    temp_filepath = os.path.join(TEMP_PATH, temp_filename)
    
    if not os.path.exists(temp_filepath):
        return jsonify({'error': 'Generated file not found'}), 404
    
    # Generate final filename with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    
    prefix_map = {
        'Neutral Letter': 'Neutral_Letters',
        'Blind Offer Letter': 'BO_Letters',
        'Offer Letter': 'Offer_Letters',
        'Multi Offer Letter': 'Multi_Offer_Letters',
        '2nd Offer Letter': '2nd_Offer_Letters',
        'Multi 2nd Offer Letter': 'Multi_2nd_Offer_Letters',
        'Contract Template': 'Contracts'
    }
    
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    cursor.execute('SELECT name FROM document_templates WHERE template_id = %s', 
                   (data.get('template_id'),))
    template_name = cursor.fetchone()['name']
    
    prefix = prefix_map.get(template_name, 'Document')
    final_filename = f"{prefix}_{timestamp}.docx"
    final_filepath = os.path.join(GENERATED_PATH, final_filename)
    
    # Move file
    os.rename(temp_filepath, final_filepath)
    
    # Save to database
    cursor.execute('''
        INSERT INTO generated_documents (p_ids, template_id, file_path, is_temp)
        VALUES (%s, %s, %s, FALSE)
    ''', (json.dumps(data.get('property_ids', [])), data.get('template_id'), 
          f'generated_documents/{final_filename}'))
    
    conn.commit()
    conn.close()
    
    # Clean up progress
    del generation_progress[progress_id]
    
    return jsonify({
        'success': True,
        'filename': final_filename,
        'download_url': f'/static/generated_documents/{final_filename}'
    })

@app.route('/api/documents/cancel', methods=['POST'])
def cancel_documents():
    """Cancel and delete temporary files."""
    data = request.json
    progress_id = data.get('progress_id')
    
    if not progress_id:
        return jsonify({'error': 'Progress ID required'}), 400
    
    progress = generation_progress.get(progress_id)
    if progress and progress.get('filename'):
        temp_filepath = os.path.join(TEMP_PATH, progress['filename'])
        if os.path.exists(temp_filepath):
            os.remove(temp_filepath)
    
    if progress_id in generation_progress:
        del generation_progress[progress_id]
    
    return jsonify({'success': True})

def generate_document_for_owner(owner, properties, template_filename, generation_date):
    """Generate document for a single owner with their properties."""
    template_path = get_template_path(template_filename)
    
    if not os.path.exists(template_path):
        raise Exception(f"Template not found: {template_filename}")
    
    doc = Document(template_path)
    
    # Prepare merge data
    merge_data = {}
    
    # Add company data
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    cursor.execute('SELECT * FROM companies WHERE c_id = 1')
    company = dict(cursor.fetchone())
    conn.close()
    
    # Flatten company data with c_ prefix
    for key, value in company.items():
        merge_data[f'c_{key[2:]}' if key.startswith('c_') else f'c_{key}'] = value or ''
    
    # Add owner data
    for key, value in owner.items():
        if value is None:
            value = ''
        merge_data[key] = value
    
    # Calculate date fields
    merge_data['m_date'] = generation_date.strftime('%B %d, %Y')
    merge_data['p_offer_accept_date'] = (generation_date + timedelta(days=15)).strftime('%B %d, %Y')
    merge_data['p_contract_expires_date'] = (generation_date + timedelta(days=180)).strftime('%B %d, %Y')
    
    # Process based on number of properties
    is_multi = len(properties) > 1
    
    if is_multi:
        # Multi-property calculations
        mo_p_price = sum(float(p.get('p_price', 0) or 0) for p in properties)
        mo_p_back_tax = sum(float(p.get('p_back_tax', 0) or 0) for p in properties)
        mo_net = mo_p_price - mo_p_back_tax
        
        merge_data['mo_p_price'] = f"{mo_p_price:.2f}"
        merge_data['mo_p_back_tax'] = f"{mo_p_back_tax:.2f}"
        merge_data['mo_net'] = f"{mo_net:.2f}"
        
        prop = properties[0]
    else:
        prop = properties[0]
        
        # Single property calculations
        p_price = float(prop.get('p_price', 0) or 0)
        p_back_tax = float(prop.get('p_back_tax', 0) or 0)
        m_net = p_price - p_back_tax
        
        merge_data['p_price'] = f"{p_price:.2f}"
        merge_data['p_back_tax'] = f"{p_back_tax:.2f}"
        merge_data['m_net'] = f"{m_net:.2f}"
    
    # Add property fields
    for key, value in prop.items():
        if not key.startswith('c_') and not key.startswith('o_') and not key.startswith('or_'):
            if value is None:
                value = ''
            merge_data[key] = value
        
    # Replace merge fields in document
    replace_merge_fields(doc, merge_data)
    
    return doc

def append_document_content(target_doc, source_doc):
    """APPEND DOCUMENT CONTENT USING PYTHON-DOCX INTERNALS."""
    source_stream = BytesIO()
    source_doc.save(source_stream)
    source_stream.seek(0)
    
    temp_doc = Document(source_stream)
    
    for child in temp_doc._element.body:
        child_xml = deepcopy(child)
        target_doc._element.body.append(child_xml)
    
def replace_merge_fields(doc, merge_data):
    """Replace placeholders by processing full paragraph text"""
    
    def merge_runs_and_replace(paragraph):
        """Merge all runs, replace placeholders, then rebuild"""
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for field_name, value in merge_data.items():
            full_text = full_text.replace(f'[{field_name}]', str(value))
        
        paragraph.clear()
        if full_text:
            paragraph.add_run(full_text)
    
    for paragraph in doc.paragraphs:
        if '[' in paragraph.text:
            merge_runs_and_replace(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[' in paragraph.text:
                        merge_runs_and_replace(paragraph)
    
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if '[' in paragraph.text:
                    merge_runs_and_replace(paragraph)
    
    for section in doc.sections:
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if '[' in paragraph.text:
                    merge_runs_and_replace(paragraph)
    
    return doc

# ============================================================================
# BULK IMPORT ROUTES
# ============================================================================

@app.route('/import')
def import_page():
    """Bulk Data Import Page"""
    return render_template('import.html')

@app.route('/api/import/template', methods=['GET'])
def download_csv_template():
    """Download a template CSV with all possible columns."""
    columns = [
        'p_status', 'p_state', 'p_county', 'p_apn', 'p_price', 'p_acres', 'p_sqft',
        'p_liens', 'p_back_tax', 'p_improvements', 'p_power', 'p_access', 'p_owned',
        'p_aquired', 'p_last_sold_date', 'p_last_sold_amount', 'p_last_transaction_date',
        'p_last_transaction_doc_type', 'p_est_value', 'p_agent_name', 'p_agent_phone',
        'p_comments', 'p_listed', 'p_address', 'p_city', 'p_terrain', 'p_short_legal',
        'p_zoning', 'p_base_tax', 'p_county_market_value', 'p_sale_price',
        'p_county_assessed_value', 'p_hoa', 'p_plat_map_link', 'p_zip',
        'p_offer_accept_date', 'p_contract_expires_date', 'p_purchased_on',
        'p_purchase_amount', 'p_purchase_closing_costs', 'p_sold_on', 'p_sold_amount',
        'p_sold_closing_costs', 'p_min_acceptable_offer', 'p_closing_company_name_purchase',
        'p_closing_company_name_sale', 'p_max_offer_amount', 'p_flood', 'p_flood_description',
        'p_use', 'p_use_code', 'p_use_description', 'p_restrictions',
        'p_waste_system_requirement', 'p_water_system_requirement', 'p_environmental',
        'p_impact_fee', 'p_betty_score', 'o_type', 'or_fname', 'or_lname', 'or_greeting',
        'o_fname', 'o_lname', 'o_2fname', 'o_2lname', 'o_3fname', 'o_3lname',
        'o_4fname', 'o_4lname', 'o_5fname', 'o_5lname', 'o_company', 'or_email',
        'or_phone', 'or_fax', 'or_m_address', 'or_m_address2', 'or_m_city',
        'or_m_state', 'or_m_zip', 'o_other_owners', 'tag_ids', 'tags', 'p_note'
    ]
    
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=columns)
    writer.writeheader()
    
    example = {
        'p_state': 'CA',
        'p_county': 'Los Angeles',
        'p_apn': '123456789',
        'o_type': 'Individual',
        'or_fname': 'John',
        'or_lname': 'Doe',
        'or_m_address': '123 Main St',
        'or_m_city': 'Los Angeles',
        'or_m_state': 'CA',
        'or_m_zip': '90210',
        'p_price': '50000',
        'p_acres': '1.5',
        'tags': 'Waterfront|Corner Lot'
    }
    writer.writerow(example)
    
    bytes_output = io.BytesIO(output.getvalue().encode('utf-8-sig'))
    bytes_output.seek(0)
    
    return send_file(
        bytes_output,
        mimetype='text/csv',
        as_attachment=True,
        download_name='import_template.csv'
    )

@app.route('/api/import/template/xlsx', methods=['GET'])
def download_xlsx_template():
    """Download an Excel template with all columns, formatted as text."""
    columns = [
        'p_status', 'p_state', 'p_county', 'p_apn', 'p_price', 'p_acres', 'p_sqft',
        'p_liens', 'p_back_tax', 'p_improvements', 'p_power', 'p_access', 'p_owned',
        'p_aquired', 'p_last_sold_date', 'p_last_sold_amount', 'p_last_transaction_date',
        'p_last_transaction_doc_type', 'p_est_value', 'p_agent_name', 'p_agent_phone',
        'p_comments', 'p_listed', 'p_address', 'p_city', 'p_terrain', 'p_short_legal',
        'p_zoning', 'p_base_tax', 'p_county_market_value', 'p_sale_price',
        'p_county_assessed_value', 'p_hoa', 'p_plat_map_link', 'p_zip',
        'p_offer_accept_date', 'p_contract_expires_date', 'p_purchased_on',
        'p_purchase_amount', 'p_purchase_closing_costs', 'p_sold_on', 'p_sold_amount',
        'p_sold_closing_costs', 'p_min_acceptable_offer', 'p_closing_company_name_purchase',
        'p_closing_company_name_sale', 'p_max_offer_amount', 'p_flood', 'p_flood_description',
        'p_use', 'p_use_code', 'p_use_description', 'p_restrictions',
        'p_waste_system_requirement', 'p_water_system_requirement', 'p_environmental',
        'p_impact_fee', 'p_betty_score', 'o_type', 'or_fname', 'or_lname', 'or_greeting',
        'o_fname', 'o_lname', 'o_2fname', 'o_2lname', 'o_3fname', 'o_3lname',
        'o_4fname', 'o_4lname', 'o_5fname', 'o_5lname', 'o_company', 'or_email',
        'or_phone', 'or_fax', 'or_m_address', 'or_m_address2', 'or_m_city',
        'or_m_state', 'or_m_zip', 'o_other_owners', 'tag_ids', 'tags', 'p_note'
    ]
    
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Import Template"
    
    # Add headers
    for col_idx, header in enumerate(columns, 1):
        cell = worksheet.cell(row=1, column=col_idx, value=header)
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    
    # Add example row
    example_data = {
        'p_state': 'CA',
        'p_county': 'Los Angeles',
        'p_apn': '320326100125',  # Long APN as string
        'o_type': 'Individual',
        'or_fname': 'John',
        'or_lname': 'Doe',
        'or_m_address': '123 Main St',
        'or_m_city': 'Los Angeles',
        'or_m_state': 'CA',
        'or_m_zip': '90210',
        'p_price': '50000',
        'p_acres': '1.5',
        'tags': 'Waterfront|Corner Lot'
    }
    
    for col_idx, header in enumerate(columns, 1):
        value = example_data.get(header, '')
        cell = worksheet.cell(row=2, column=col_idx, value=value)
        # Ensure text format for all cells (prevents scientific notation)
        cell.number_format = '@'
    
    # Auto-adjust column widths
    for col_idx, header in enumerate(columns, 1):
        column_letter = get_column_letter(col_idx)
        worksheet.column_dimensions[column_letter].width = max(len(header) + 2, 15)
    
    # Save to BytesIO
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='import_template.xlsx'
    )

@app.route('/api/import', methods=['POST'])
def process_import():
    """Process CSV or XLSX import with all-or-nothing transaction and detailed audit logging."""
    conn = None
    audit_log = []
    
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'}), 400
        
        mode = request.form.get('mode', 'new')
        dont_change_status = request.form.get('dont_change_status', 'false').lower() == 'true'
        
        # Determine file type and parse accordingly
        filename = file.filename.lower()
        rows = []
        
        if filename.endswith('.xlsx'):
            # Handle Excel file
            audit_log.append("File type: Excel (.xlsx)")
            workbook = openpyxl.load_workbook(file.stream, data_only=True)
            worksheet = workbook.active
            
            # Get headers from first row
            headers = []
            first_row = next(worksheet.iter_rows(values_only=True))
            headers = [str(cell).strip().lower() if cell else '' for cell in first_row]
            
            # Convert remaining rows to dict format (same as CSV DictReader)
            for row in worksheet.iter_rows(min_row=2, values_only=True):
                if not any(cell is not None and str(cell).strip() for cell in row):
                    continue  # Skip empty rows
                
                row_dict = {}
                for idx, header in enumerate(headers):
                    if idx < len(row):
                        cell_value = row[idx]
                        # CRITICAL: Convert all values to string to prevent scientific notation
                        if cell_value is None:
                            row_dict[header] = ''
                        else:
                            # Force string conversion - this prevents 3.20326E+11 formatting
                            row_dict[header] = str(cell_value).strip()
                    else:
                        row_dict[header] = ''
                rows.append(row_dict)
                
        elif filename.endswith('.csv'):
            # Handle CSV file (existing logic)
            audit_log.append("File type: CSV")
            file_content = file.stream.read()
            
            try:
                stream = io.StringIO(file_content.decode('utf-8-sig'), newline=None)
                audit_log.append("File encoding: UTF-8")
            except UnicodeDecodeError:
                try:
                    stream = io.StringIO(file_content.decode('latin-1'), newline=None)
                    audit_log.append("File encoding: Latin-1")
                except UnicodeDecodeError:
                    stream = io.StringIO(file_content.decode('cp1252'), newline=None)
                    audit_log.append("File encoding: Windows-1252")
            
            reader = csv.DictReader(stream)
            rows = list(reader)
        else:
            return jsonify({'success': False, 'message': 'Unsupported file type. Please upload .csv or .xlsx files only.'}), 400

        if not rows:
            return jsonify({'success': False, 'message': 'File is empty or contains no data rows'}), 400
        
        # Normalize headers (lowercase, strip whitespace)
        if rows:
            first_row = rows[0]
            normalized_rows = []
            for row in rows:
                normalized_row = {}
                for key, value in row.items():
                    if key:
                        normalized_key = key.lower().strip()
                        normalized_row[normalized_key] = value
                normalized_rows.append(normalized_row)
            rows = normalized_rows
        
        headers = list(rows[0].keys()) if rows else []
        
        # Continue with existing validation logic...
        mandatory_fields = ['p_state', 'p_county', 'p_apn', 'o_type', 'or_m_address', 
                           'or_m_city', 'or_m_state', 'or_m_zip']
        
        missing_fields = [f for f in mandatory_fields if f.lower() not in headers]
        if missing_fields:
            return jsonify({
                'success': False, 
                'message': f'Missing mandatory columns: {", ".join(missing_fields)}'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Filter out completely empty rows (existing logic)
        valid_rows = []
        for idx, row in enumerate(rows, start=2):
            has_data = any(str(row.get(field, '')).strip() for field in mandatory_fields)
            if has_data:
                valid_rows.append((idx, row))
            else:
                audit_log.append(f"Row {idx}: SKIPPED - Empty row")
        
        # Check for duplicate APNs within the file (existing logic)
        apn_locations = {}
        duplicate_errors = []
        
        for idx, row in valid_rows:
            p_apn = str(row.get('p_apn', '')).strip()
            p_county = str(row.get('p_county', '')).strip()
            p_state = str(row.get('p_state', '')).strip()
            
            prop_key = (p_apn.lower(), p_county.lower(), p_state.lower())
            
            if prop_key in apn_locations:
                first_row = apn_locations[prop_key]
                duplicate_errors.append(f"Row {idx}: Duplicate APN '{p_apn}' - duplicate of Row {first_row}")
            else:
                apn_locations[prop_key] = idx
        
        if duplicate_errors:
            audit_log.append(f"DUPLICATE APN DETECTION FAILED")
            return jsonify({
                'success': False,
                'message': f'Duplicate APNs found within import file.',
                'errors': duplicate_errors,
                'audit': audit_log
            }), 400
        
        # Rest of existing validation and import logic continues unchanged...
        errors = []
        for idx, row in valid_rows:
            row_errors = validate_import_row(row, idx, headers)
            if row_errors:
                errors.extend(row_errors)
                audit_log.append(f"Row {idx}: VALIDATION FAILED - {row_errors}")
        
        if errors:
            return jsonify({
                'success': False,
                'message': f'{len(errors)} validation errors found.',
                'errors': errors,
                'audit': audit_log
            }), 400
        
        # Group by owner address (existing logic continues...)
        owner_groups = {}
        for row_num, row in valid_rows:
            owner_key = (
                str(row.get('or_m_address', '')).strip().lower(),
                str(row.get('or_m_city', '')).strip().lower(),
                str(row.get('or_m_state', '')).strip().lower()
            )
            
            if owner_key not in owner_groups:
                owner_groups[owner_key] = {
                    'owner_data': row,
                    'properties': []
                }
            owner_groups[owner_key]['properties'].append((row_num, row))
        
        stats = {
            'new_owners': 0,
            'existing_owners': 0,
            'new_properties': 0,
            'existing_properties': 0
        }
        
        import_operations = []
        
        for owner_key, group in owner_groups.items():
            # Check for existing owner (existing logic)
            cursor.execute(
                """SELECT or_id, or_fname, or_lname, or_m_address 
                   FROM owners 
                   WHERE TRIM(LOWER(or_m_address)) = %s 
                   AND TRIM(LOWER(or_m_city)) = %s 
                   AND TRIM(LOWER(or_m_state)) = %s""",
                owner_key
            )
            owner_row = cursor.fetchone()
            
            if owner_row:
                or_id = owner_row['or_id']
                is_new_owner = False
                stats['existing_owners'] += 1
                audit_log.append(f"Owner {owner_key}: FOUND EXISTING (ID: {or_id})")
            else:
                or_id = None
                is_new_owner = True
                stats['new_owners'] += 1
                audit_log.append(f"Owner {owner_key}: NEW OWNER")
            
            for row_num, row in group['properties']:
                p_apn = str(row.get('p_apn', '')).strip()
                p_county = str(row.get('p_county', '')).strip()
                p_state = str(row.get('p_state', '')).strip()
                
                # Check for existing property (existing logic)
                cursor.execute(
                    """SELECT p_id, p_status_id, p_apn 
                       FROM properties 
                       WHERE TRIM(p_apn) = %s AND TRIM(p_county) = %s AND TRIM(p_state) = %s""",
                    (p_apn, p_county, p_state)
                )
                prop_row = cursor.fetchone()
                
                if prop_row:
                    existing_p_id = prop_row['p_id']
                    existing_status_id = prop_row['p_status_id']
                    is_new_property = False
                    audit_log.append(f"Row {row_num} (APN {p_apn}): PROPERTY EXISTS")
                else:
                    existing_p_id = None
                    existing_status_id = None
                    is_new_property = True
                    audit_log.append(f"Row {row_num} (APN {p_apn}): NEW PROPERTY")
                
                if mode == 'new':
                    if not is_new_property:
                        stats['existing_properties'] += 1
                        audit_log.append(f"Row {row_num}: SKIPPED (Mode=New)")
                        continue
                    action = 'create'
                    stats['new_properties'] += 1 
                else:
                    if is_new_property:
                        action = 'create'
                        stats['new_properties'] += 1
                    else:
                        action = 'update'
                        stats['existing_properties'] += 1
                
                import_operations.append({
                    'row_num': row_num,
                    'row': row,
                    'action': action,
                    'is_new_owner': is_new_owner,
                    'or_id': or_id,
                    'existing_p_id': existing_p_id,
                    'existing_status_id': existing_status_id,
                    'apn': p_apn
                })
        
        if not import_operations:
            conn.close()
            return jsonify({
                'success': True,
                'message': 'No new records to import.',
                'audit': audit_log
            })
        
        timestamp = datetime.now().strftime('%b %d, %Y, %I:%M %p')
        created_owners = {}
        
        for op in import_operations:
            row = op['row']
            
            try:
                if op['is_new_owner']:
                    owner_key = (
                        str(row.get('or_m_address', '')).strip().lower(),
                        str(row.get('or_m_city', '')).strip().lower(),
                        str(row.get('or_m_state', '')).strip().lower()
                    )
                    
                    if owner_key in created_owners:
                        current_or_id = created_owners[owner_key]
                        audit_log.append(f"Row {op['row_num']}: Using cached owner ID {current_or_id}")
                    else:
                        owner_data = prepare_owner_data_cleaned(row)
                        cursor.execute(
                            """INSERT INTO owners (o_type, or_fname, or_lname, or_email, or_phone, or_fax,
                                o_fname, o_lname, o_2fname, o_2lname, o_3fname, o_3lname,
                                o_4fname, o_4lname, o_5fname, o_5lname, o_company,
                                or_m_address, or_m_address2, or_m_city, or_m_state, or_m_zip,
                                o_other_owners)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            RETURNING or_id""",
                            owner_data
                        )
                        current_or_id = cursor.fetchone()['or_id']
                        created_owners[owner_key] = current_or_id
                        audit_log.append(f"Row {op['row_num']}: Created owner ID {current_or_id}")
                else:
                    current_or_id = op['or_id']
                    audit_log.append(f"Row {op['row_num']}: Using existing owner ID {current_or_id}")
                
                csv_status = str(row.get('p_status', '')).strip()
                
                if op['action'] == 'create':
                    if not csv_status:
                        csv_status = 'Prospect'
                    
                    cursor.execute('SELECT status_id FROM statuses WHERE s_status = %s', (csv_status,))
                    status_row = cursor.fetchone()
                    if status_row:
                        status_id = status_row['status_id']
                    else:
                        cursor.execute('SELECT status_id FROM statuses WHERE s_status = %s', ('Prospect',))
                        prospect_row = cursor.fetchone()
                        status_id = prospect_row['status_id'] if prospect_row else 1
                        if csv_status != 'Prospect':
                            audit_log.append(f"Row {op['row_num']}: Warning - Status '{csv_status}' not found")
                    
                    prop_data = prepare_property_data(row, current_or_id, status_id, timestamp)
                    
                    if len(prop_data) != 69:
                        error_msg = f"Row {op['row_num']}: DATA LENGTH ERROR"
                        audit_log.append(error_msg)
                        raise ValueError(error_msg)
                    
                    cursor.execute(
                        """INSERT INTO properties 
                        (or_id, p_status_id, p_apn, p_county, p_state, p_longstate, p_short_legal,
                         p_address, p_city, p_zip, p_acres, p_sqft, p_terrain, p_zoning, p_use,
                         p_use_code, p_use_description, p_impact_fee, p_environmental, p_restrictions,
                         p_waste_system_requirement, p_water_system_requirement, p_survey, p_flood,
                         p_flood_description, p_base_tax, p_hoa, p_liens, p_back_tax,
                         p_county_assessed_value, p_county_market_value, p_comp_market_value,
                         p_improvements, p_power, p_access, p_owned, p_aquired, p_est_value,
                         p_min_acceptable_offer, p_price, p_max_offer_amount, p_m_date,
                         p_offer_accept_date, p_contract_expires_date, p_comments, p_note,
                         p_plat_map_link, p_viable, p_betty_score, p_create_time, p_last_updated,
                         p_status_last_updated, p_listed, p_agent_name, p_agent_phone, p_purchased_on,
                         p_purchase_amount, p_purchase_closing_costs, p_closing_company_name_purchase,
                         p_sold_on, p_sold_amount, p_buyer, p_sold_closing_costs, p_profit,
                         p_closing_company_name_sale, p_last_sold_date, p_last_sold_amount,
                         p_last_transaction_date, p_last_transaction_doc_type)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 
                                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING p_id""",
                        prop_data
                    )
                    p_id = cursor.fetchone()['p_id']
                    audit_log.append(f"Row {op['row_num']}: Created property ID {p_id}")
                    handle_tags_import(cursor, row.get('tags', ''), p_id)
                    
                else:
                    updates = []
                    params = []
                    
                    cursor.execute('''
                        SELECT s.s_status 
                        FROM properties p 
                        LEFT JOIN statuses s ON p.p_status_id = s.status_id 
                        WHERE p.p_id = %s
                    ''', (op['existing_p_id'],))
                    current_status_row = cursor.fetchone()
                    current_status_name = current_status_row['s_status'] if current_status_row else ''
                    
                    updatable_fields = [
                        'p_price', 'p_acres', 'p_sqft', 'p_liens', 'p_back_tax', 'p_improvements',
                        'p_power', 'p_access', 'p_owned', 'p_aquired', 'p_last_sold_date',
                        'p_last_sold_amount', 'p_last_transaction_date', 'p_last_transaction_doc_type',
                        'p_est_value', 'p_agent_name', 'p_agent_phone', 'p_comments', 'p_listed',
                        'p_address', 'p_city', 'p_terrain', 'p_short_legal', 'p_zoning', 'p_base_tax',
                        'p_county_market_value', 'p_sale_price', 'p_county_assessed_value', 'p_hoa',
                        'p_plat_map_link', 'p_zip', 'p_offer_accept_date', 'p_contract_expires_date',
                        'p_purchased_on', 'p_purchase_amount', 'p_purchase_closing_costs',
                        'p_closing_company_name_purchase', 'p_sold_on', 'p_sold_amount',
                        'p_sold_closing_costs', 'p_min_acceptable_offer', 'p_closing_company_name_sale',
                        'p_max_offer_amount', 'p_flood', 'p_flood_description', 'p_use', 'p_use_code',
                        'p_use_description', 'p_restrictions', 'p_waste_system_requirement',
                        'p_water_system_requirement', 'p_environmental', 'p_impact_fee', 'p_betty_score',
                        'p_note'
                    ]
                    
                    for field in updatable_fields:
                        if field in row and str(row[field]).strip():
                            cursor.execute(f'SELECT {field} FROM properties WHERE p_id = %s', (op['existing_p_id'],))
                            current = cursor.fetchone()[field]
                            if current is None or str(current).strip() == '':
                                updates.append(f"{field} = %s")
                                params.append(row[field])
                    
                    if not dont_change_status and current_status_name != 'FILE CLOSED':
                        if not csv_status:
                            csv_status = 'Prospect'
                        
                        cursor.execute('SELECT status_id FROM statuses WHERE s_status = %s', (csv_status,))
                        new_status_row = cursor.fetchone()
                        
                        if new_status_row:
                            new_status_id = new_status_row['status_id']
                            updates.append("p_status_id = %s")
                            params.append(new_status_id)
                            updates.append("p_status_last_updated = %s")
                            params.append(timestamp)
                            audit_log.append(f"Row {op['row_num']}: Updated status")
                        else:
                            audit_log.append(f"Row {op['row_num']}: Warning - Status not found")
                    else:
                        if dont_change_status:
                            audit_log.append(f"Row {op['row_num']}: Status not changed")
                        elif current_status_name == 'FILE CLOSED':
                            audit_log.append(f"Row {op['row_num']}: Status not changed (FILE CLOSED)")
                    
                    if updates:
                        updates.append("p_last_updated = %s")
                        params.append(timestamp)
                        params.append(op['existing_p_id'])
                        
                        sql = f"UPDATE properties SET {', '.join(updates)} WHERE p_id = %s"
                        cursor.execute(sql, params)
                        mark_record_modified('properties', op['existing_p_id'], conn)
                        audit_log.append(f"Row {op['row_num']}: Updated property")
                    
                    handle_tags_import(cursor, row.get('tags', ''), op['existing_p_id'], append=True)
                    
            except Exception as row_error:
                error_msg = f"Row {op['row_num']} (APN {op['apn']}): ERROR - {str(row_error)}"
                audit_log.append(error_msg)
                raise
        
        conn.commit()

        # Mark imported records for sync
        imported_ids = [op.get('p_id') for op in import_operations if op.get('p_id')]
        if imported_ids:
            placeholders = ','.join(['%s'] * len(imported_ids))
            cursor.execute(f"""
                UPDATE properties 
                SET sync_status = 'pending', modified_at = NOW() 
                WHERE p_id IN ({placeholders})
            """, imported_ids)
        
        return jsonify({
            'success': True,
            'message': f'{len(import_operations)} rows imported successfully!',
            'new_owners': stats['new_owners'],
            'existing_owners': stats['existing_owners'],
            'new_properties': stats['new_properties'],
            'existing_properties': stats['existing_properties'],
            'audit': audit_log
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'Import failed: ' + str(e),
            'errors': [str(e)],
            'audit': audit_log
        }), 500
    finally:
        if conn:
            conn.close()

def prepare_owner_data_cleaned(row):
    """Extract owner fields with whitespace stripped."""
    return (
        str(row.get('o_type', 'Individual')).strip(),
        str(row.get('or_fname', '')).strip(),
        str(row.get('or_lname', '')).strip(),
        str(row.get('or_email', '')).strip(),
        str(row.get('or_phone', '')).strip(),
        str(row.get('or_fax', '')).strip(),
        str(row.get('o_fname', '')).strip(),
        str(row.get('o_lname', '')).strip(),
        str(row.get('o_2fname', '')).strip(),
        str(row.get('o_2lname', '')).strip(),
        str(row.get('o_3fname', '')).strip(),
        str(row.get('o_3lname', '')).strip(),
        str(row.get('o_4fname', '')).strip(),
        str(row.get('o_4lname', '')).strip(),
        str(row.get('o_5fname', '')).strip(),
        str(row.get('o_5lname', '')).strip(),
        str(row.get('o_company', '')).strip(),
        str(row.get('or_m_address', '')).strip(),
        str(row.get('or_m_address2', '')).strip(),
        str(row.get('or_m_city', '')).strip(),
        str(row.get('or_m_state', '')).strip(),
        str(row.get('or_m_zip', '')).strip(),
        True if any([str(row.get(f'o_{i}fname', '')).strip() or str(row.get(f'o_{i}lname', '')).strip() for i in range(2,6)]) else False
    )

def validate_import_row(row, row_num, headers):
    """Validate a single row and return list of error messages."""
    errors = []
    
    mandatory = {
        'p_state': 'p_state',
        'p_county': 'p_county', 
        'p_apn': 'p_apn',
        'o_type': 'o_type',
        'or_m_address': 'or_m_address',
        'or_m_city': 'or_m_city',
        'or_m_state': 'or_m_state',
        'or_m_zip': 'or_m_zip'
    }
    
    for field, header in mandatory.items():
        value = row.get(header, '').strip()
        if not value:
            errors.append(f"Row {row_num}: {header} is mandatory")
    
    o_type = row.get('o_type', '').strip()
    if o_type and o_type not in ['Individual', 'Company']:
        errors.append(f"Row {row_num}: o_type must be 'Individual' or 'Company'")
    
    if o_type == 'Individual':
        if not row.get('or_fname', '').strip():
            errors.append(f"Row {row_num}: or_fname required for Individual")
        if not row.get('or_lname', '').strip():
            errors.append(f"Row {row_num}: or_lname required for Individual")
    elif o_type == 'Company':
        if not row.get('o_company', '').strip():
            errors.append(f"Row {row_num}: o_company required for Company")
    
    return errors

def prepare_property_data(row, or_id, status_id, timestamp):
    """Extract and calculate property fields from row."""
    p_state = row.get('p_state', '')
    
    p_acres = clean_numeric_field(row.get('p_acres'))
    p_sqft = clean_numeric_field(row.get('p_sqft'))
    
    if p_acres and not p_sqft:
        try:
            p_sqft = float(p_acres) * 43560
        except:
            p_sqft = None
    elif p_sqft and not p_acres:
        try:
            p_acres = float(p_sqft) / 43560
        except:
            p_acres = None
    
    def convert_date(date_str):
        if not date_str:
            return None
        date_str = str(date_str).strip()
        if not date_str:
            return None
        try:
            # Try MM/DD/YYYY first (import format)
            dt = datetime.strptime(date_str, '%m/%d/%Y')
            return dt.strftime('%Y-%m-%d')
        except ValueError:
            try:
                # Try YYYY-MM-DD (already correct format)
                dt = datetime.strptime(date_str, '%Y-%m-%d')
                return date_str
            except ValueError:
                try:
                    # Try "Dec 18, 2018" format (create_time format)
                    dt = datetime.strptime(date_str, '%b %d, %Y')
                    return dt.strftime('%Y-%m-%d')
                except ValueError:
                    # Return as-is if unparseable
                    return date_str
    
    p_create_time = str(row.get('p_create_time', '')).strip() or timestamp
    p_last_updated = str(row.get('p_last_updated', '')).strip() or timestamp
    p_status_last_updated = str(row.get('p_status_last_updated', '')).strip() or timestamp
    
    return (
        or_id, status_id, row.get('p_apn', ''), row.get('p_county', ''), p_state,
        get_longstate_from_abbr(p_state), row.get('p_short_legal', ''),
        row.get('p_address', ''), row.get('p_city', ''), row.get('p_zip', ''),
        p_acres, p_sqft, row.get('p_terrain', ''), row.get('p_zoning', ''),
        row.get('p_use', ''), row.get('p_use_code', ''),
        row.get('p_use_description', ''),
        clean_numeric_field(row.get('p_impact_fee')),
        row.get('p_environmental', ''), row.get('p_restrictions', ''),
        row.get('p_waste_system_requirement', ''),
        row.get('p_water_system_requirement', ''),
        True if row.get('p_survey') in ['1', 'Yes', 'True', 'yes'] else False,
        row.get('p_flood', ''), row.get('p_flood_description', ''),
        clean_numeric_field(row.get('p_base_tax')),
        clean_numeric_field(row.get('p_hoa')),
        clean_numeric_field(row.get('p_liens')),
        clean_numeric_field(row.get('p_back_tax')),
        clean_numeric_field(row.get('p_county_assessed_value')),
        clean_numeric_field(row.get('p_county_market_value')),
        clean_numeric_field(row.get('p_comp_market_value')),
        row.get('p_improvements', ''),
        row.get('p_power', ''), row.get('p_access', ''),
        row.get('p_owned', ''), row.get('p_aquired', ''),
        clean_numeric_field(row.get('p_est_value')),
        clean_numeric_field(row.get('p_min_acceptable_offer')),
        clean_numeric_field(row.get('p_price')),
        clean_numeric_field(row.get('p_max_offer_amount')),
        convert_date(row.get('p_m_date')),
        convert_date(row.get('p_offer_accept_date')),
        convert_date(row.get('p_contract_expires_date')),
        row.get('p_comments', ''), row.get('p_note', ''),
        row.get('p_plat_map_link', ''),
        True if row.get('p_viable') in ['1', 'Yes', 'True', 'yes'] else False,
        clean_numeric_field(row.get('p_betty_score')),
        p_create_time, p_last_updated, p_status_last_updated,
        True if row.get('p_listed') in ['1', 'Yes', 'True', 'yes'] else False,
        row.get('p_agent_name', ''), row.get('p_agent_phone', ''),
        convert_date(row.get('p_purchased_on')),
        clean_numeric_field(row.get('p_purchase_amount')),
        clean_numeric_field(row.get('p_purchase_closing_costs')),
        row.get('p_closing_company_name_purchase', ''),
        convert_date(row.get('p_sold_on')),
        clean_numeric_field(row.get('p_sold_amount')),
        row.get('p_buyer', ''),
        clean_numeric_field(row.get('p_sold_closing_costs')),
        clean_numeric_field(row.get('p_profit')),
        row.get('p_closing_company_name_sale', ''),
        convert_date(row.get('p_last_sold_date')),
        clean_numeric_field(row.get('p_last_sold_amount')),
        convert_date(row.get('p_last_transaction_date')),
        row.get('p_last_transaction_doc_type', '')
    )

def handle_tags_import(cursor, tags_str, p_id, append=False):
    """Handle tag import with pipe delimiter."""
    if not tags_str:
        return
    
    tags = [t.strip() for t in tags_str.split('|') if t.strip()]
    
    for tag_name in tags:
        cursor.execute('SELECT tag_id FROM tags WHERE tag_name = %s', (tag_name,))
        tag_row = cursor.fetchone()
        
        if tag_row:
            tag_id = tag_row['tag_id']
        else:
            cursor.execute('INSERT INTO tags (tag_name, tag_color) VALUES (%s, %s) RETURNING tag_id',
                         (tag_name, '#808080'))
            tag_id = cursor.fetchone()['tag_id']
        
        if append:
            cursor.execute('SELECT 1 FROM property_tags WHERE p_id = %s AND tag_id = %s', (p_id, tag_id))
            if cursor.fetchone():
                continue
        
        cursor.execute('INSERT INTO property_tags (p_id, tag_id) VALUES (%s, %s)', (p_id, tag_id))

@app.route('/api/properties/bulk-delete', methods=['POST'])
def bulk_delete_properties():
    """Delete multiple property records and propagate to cloud."""
    conn = None
    try:
        data = request.json
        p_ids = data.get('p_ids', [])
        
        if not p_ids:
            return jsonify({'error': 'No property IDs provided'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 1. Insert tombstones for all deletions
        for p_id in p_ids:
            cursor.execute('''
                INSERT INTO sync_deletions (table_name, record_id, sync_status)
                VALUES (%s, %s, %s)
                ON CONFLICT (table_name, record_id) DO UPDATE SET
                deleted_at = NOW(), sync_status = 'pending', cloud_deleted = FALSE
            ''', ('properties', p_id, 'pending'))
        
        # 2. Try to delete from Supabase immediately
        try:
            from sync_service import ValiantLandSync
            sync = ValiantLandSync()
            for p_id in p_ids:
                try:
                    sync.supabase.table('properties').delete().eq('p_id', p_id).execute()
                    cursor.execute('UPDATE sync_deletions SET cloud_deleted = TRUE WHERE table_name = %s AND record_id = %s', ('properties', p_id))
                except Exception as single_error:
                    print(f"Cloud delete queued for property {p_id}: {single_error}")
        except Exception as cloud_error:
            print(f"Bulk cloud delete error (will retry later): {cloud_error}")
        
        # 3. Delete related file tracking
        for p_id in p_ids:
            cursor.execute('DELETE FROM file_sync WHERE local_path LIKE %s', (f'%/p_{p_id}/%',))
        
        # 4. Delete local properties
        placeholders = ','.join(['%s'] * len(p_ids))
        cursor.execute(f'DELETE FROM properties WHERE p_id IN ({placeholders})', p_ids)
        
        deleted_count = cursor.rowcount
        conn.commit()
        
        return jsonify({
            'success': True, 
            'deleted_count': deleted_count,
            'message': f'{deleted_count} records deleted and queued for cloud sync'
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"ERROR bulk deleting: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

# ============================================================================
# ADVANCED SEARCH ROUTES
# ============================================================================

@app.route('/api/search', methods=['POST'])
def advanced_search():
    """Advanced search with multiple criteria and AND/OR logic."""
    conn = None
    try:
        data = request.json
        logic = data.get('logic', 'AND')
        rules = data.get('rules', [])
        
        if not rules:
            return jsonify({'error': 'No search criteria provided'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        query_parts = []
        params = []
        
        number_fields = {
            'p_base_tax', 'p_back_tax', 'p_betty_score', 'p_comp_market_value', 
            'p_county_assessed_value', 'p_county_market_value', 'p_hoa', 'p_impact_fee',
            'p_liens', 'p_max_offer_amount', 'p_min_acceptable_offer', 'p_price',
            'p_est_value', 'p_acres', 'p_sqft', 'p_purchase_amount', 
            'p_purchase_closing_costs', 'p_sold_amount', 'p_sold_closing_costs', 
            'p_sale_price', 'p_id', 'or_id', 'p_owned'
        }
        
        date_fields = {
            'p_contract_expires_date', 'p_create_time', 'p_offer_accept_date',
            'p_purchased_on', 'p_sold_on', 'p_last_sold_date', 'p_last_transaction_date',
            'p_status_last_updated', 'p_m_date', 'p_last_updated'
        }

        simple_date_fields = {
            'p_contract_expires_date', 'p_offer_accept_date', 'p_m_date',
            'p_purchased_on', 'p_sold_on', 'p_last_sold_date', 'p_last_transaction_date'
        }
        
        boolean_fields = {
            'p_viable', 'p_listed', 'p_survey', 'p_flood', 'p_improvements', 
            'p_power', 'p_access'
        }
        
        text_fields = {
            'o_2fname', 'o_2lname', 'p_apn', 'p_comments', 'o_company', 'p_county',
            'or_fname', 'p_flood_description', 'or_lname',
            'or_m_address', 'or_m_address2', 'or_m_city', 'or_m_state', 'or_m_zip',
            'p_note', 'o_fname', 'o_lname', 'o_type', 'or_phone', 'p_address',
            'p_city', 'p_state', 'p_zip', 'p_closing_company_name_purchase',
            'p_closing_company_name_sale', 'p_agent_name', 'p_agent_phone',
            'p_restrictions', 'p_short_legal', 'p_status', 'p_terrain',
            'p_aquired', 'p_use', 'p_use_code', 'p_use_description', 'p_zoning',
            'or_email', 'p_waste_system_requirement', 'p_water_system_requirement',
            'p_environmental'
        }
        
        needs_owners = False
        needs_tags = False
        needs_statuses = False
        
        for rule in rules:
            field = rule.get('field')
            operator = rule.get('operator')
            value = rule.get('value', '')
            value2 = rule.get('value2', '')
            
            if field == '-1' or not field:
                continue
                
            if field == 'all_fields':
                needs_owners = True
                needs_statuses = True
                all_field_conditions = []
                for text_field in text_fields:
                    if text_field == 'p_status':
                        table_field = "s.s_status"
                    elif text_field.startswith('or_') or text_field.startswith('o_'):
                        table_field = f"o.{text_field}"
                    else:
                        table_field = f"p.{text_field}"
                    
                    if operator == 'contains':
                        all_field_conditions.append(f"{table_field} LIKE %s")
                        params.append(f'%{value}%')
                    elif operator == 'not_contains':
                        all_field_conditions.append(f"{table_field} NOT LIKE %s")
                        params.append(f'%{value}%')
                    elif operator == 'equal':
                        all_field_conditions.append(f"{table_field} = %s")
                        params.append(value)
                    elif operator == 'not_equal':
                        all_field_conditions.append(f"{table_field} != %s")
                        params.append(value)
                    elif operator == 'empty':
                        all_field_conditions.append(f"({table_field} IS NULL OR {table_field} = '')")
                    elif operator == 'not_empty':
                        all_field_conditions.append(f"({table_field} IS NOT NULL AND {table_field} != '')")
                
                if all_field_conditions:
                    query_parts.append(f"({' OR '.join(all_field_conditions)})")
                continue
            
            if field.startswith('or_') or field.startswith('o_'):
                table_field = f"o.{field}"
                needs_owners = True
            elif field in ('tag_ids', 'tags'):
                tag_field = 'tag_id' if field == 'tag_ids' else 'tag_name'

                if operator == 'equal':
                    condition = f"EXISTS (SELECT 1 FROM property_tags pt_sub JOIN tags t_sub ON pt_sub.tag_id = t_sub.tag_id WHERE pt_sub.p_id = p.p_id AND t_sub.{tag_field} = %s)"
                    params.append(value)
                elif operator == 'not_equal':
                    condition = f"NOT EXISTS (SELECT 1 FROM property_tags pt_sub JOIN tags t_sub ON pt_sub.tag_id = t_sub.tag_id WHERE pt_sub.p_id = p.p_id AND t_sub.{tag_field} = %s)"
                    params.append(value)
                elif operator == 'contains':
                    condition = f"EXISTS (SELECT 1 FROM property_tags pt_sub JOIN tags t_sub ON pt_sub.tag_id = t_sub.tag_id WHERE pt_sub.p_id = p.p_id AND t_sub.{tag_field} LIKE %s)"
                    params.append(f'%' + value + '%')
                elif operator == 'not_contains':
                    condition = f"NOT EXISTS (SELECT 1 FROM property_tags pt_sub JOIN tags t_sub ON pt_sub.tag_id = t_sub.tag_id WHERE pt_sub.p_id = p.p_id AND t_sub.{tag_field} LIKE %s)"
                    params.append(f'%' + value + '%')
                elif operator == 'empty':
                    condition = f"NOT EXISTS (SELECT 1 FROM property_tags pt_sub WHERE pt_sub.p_id = p.p_id)"
                elif operator == 'not_empty':
                    condition = f"EXISTS (SELECT 1 FROM property_tags pt_sub WHERE pt_sub.p_id = p.p_id)"

                if condition:
                    query_parts.append(condition)
                continue
            elif field == 'p_status':
                table_field = "s.s_status"
                needs_statuses = True
            else:
                table_field = f"p.{field}"
            
            condition = None
            
            if field in number_fields:
                if operator == 'equal':
                    condition = f"{table_field} = %s"
                    params.append(value)
                elif operator == 'not_equal':
                    condition = f"{table_field} != %s"
                    params.append(value)
                elif operator == 'less':
                    condition = f"{table_field} < %s"
                    params.append(value)
                elif operator == 'less_equal':
                    condition = f"{table_field} <= %s"
                    params.append(value)
                elif operator == 'greater':
                    condition = f"{table_field} > %s"
                    params.append(value)
                elif operator == 'greater_equal':
                    condition = f"{table_field} >= %s"
                    params.append(value)
                elif operator == 'between':
                    condition = f"{table_field} BETWEEN %s AND %s"
                    params.append(value)
                    params.append(value2)
                elif operator == 'not_between':
                    condition = f"{table_field} NOT BETWEEN %s AND %s"
                    params.append(value)
                    params.append(value2)
                elif operator == 'is_null':
                    condition = f"{table_field} IS NULL"
                elif operator == 'is_not_null':
                    condition = f"{table_field} IS NOT NULL"
                    
            elif field in date_fields:
                if field in ('p_create_time', 'p_last_updated', 'p_status_last_updated'):
                    # Handle mixed timestamp formats: "Dec 18, 2018, 7:38 pm" vs "02/14/2026 02:30 PM"
                    if operator == 'is_not_null':
                        condition = f"{table_field} IS NOT NULL AND {table_field} != ''"
                    elif operator == 'is_null':
                        condition = f"{table_field} IS NULL OR {table_field} = ''"
                    elif operator in ('less', 'less_equal', 'greater', 'greater_equal'):
                        op_map = {'less': '<', 'less_equal': '<=', 'greater': '>', 'greater_equal': '>='}
                        sql_op = op_map[operator]
                        # Extract date from timestamp, handling both formats
                        condition = f"""CASE 
                            WHEN {table_field} ~ '^[A-Za-z]' 
                                THEN TO_TIMESTAMP({table_field}, 'Mon DD, YYYY, HH12:MI am')::date 
                            ELSE TO_TIMESTAMP({table_field}, 'MM/DD/YYYY HH12:MI AM')::date 
                        END {sql_op} TO_DATE(%s, 'YYYY-MM-DD')"""
                        params.append(value)
                    elif operator == 'between':
                        condition = f"""CASE 
                            WHEN {table_field} ~ '^[A-Za-z]' 
                                THEN TO_TIMESTAMP({table_field}, 'Mon DD, YYYY, HH12:MI am')::date 
                            ELSE TO_TIMESTAMP({table_field}, 'MM/DD/YYYY HH12:MI AM')::date 
                        END BETWEEN TO_DATE(%s, 'YYYY-MM-DD') AND TO_DATE(%s, 'YYYY-MM-DD')"""
                        params.append(value)
                        params.append(value2)
                    elif operator == 'not_between':
                        condition = f"""CASE 
                            WHEN {table_field} ~ '^[A-Za-z]' 
                                THEN TO_TIMESTAMP({table_field}, 'Mon DD, YYYY, HH12:MI am')::date 
                            ELSE TO_TIMESTAMP({table_field}, 'MM/DD/YYYY HH12:MI AM')::date 
                        END NOT BETWEEN TO_DATE(%s, 'YYYY-MM-DD') AND TO_DATE(%s, 'YYYY-MM-DD')"""
                        params.append(value)
                        params.append(value2)
                    else:
                        condition = f"{table_field} = %s"
                        params.append(value)
                else:
                    # Regular date fields (YYYY-MM-DD)
                    if operator == 'is_not_null':
                        condition = f"{table_field} IS NOT NULL AND {table_field} != ''"
                    elif operator == 'is_null':
                        condition = f"{table_field} IS NULL OR {table_field} = ''"
                    elif operator in ('less', 'less_equal', 'greater', 'greater_equal'):
                        op_map = {'less': '<', 'less_equal': '<=', 'greater': '>', 'greater_equal': '>='}
                        sql_op = op_map[operator]
                        condition = f"{table_field} {sql_op} %s"
                        params.append(value)
                    elif operator == 'between':
                        condition = f"{table_field} BETWEEN %s AND %s"
                        params.append(value)
                        params.append(value2)
                    elif operator == 'not_between':
                        condition = f"{table_field} NOT BETWEEN %s AND %s"
                        params.append(value)
                        params.append(value2)
                    else:
                        condition = f"{table_field} = %s"
                        params.append(value)
                
                if condition:
                    query_parts.append(condition)
                continue

                    
            elif field in boolean_fields:
                bool_value = True if value.lower() in ('yes', 'true', '1') else False
                if operator == 'equal':
                    condition = f"{table_field} = %s"
                    params.append(bool_value)
                elif operator == 'not_equal':
                    condition = f"{table_field} != %s"
                    params.append(bool_value)
                    
            else:
                if operator == 'equal':
                    condition = f"{table_field} = %s"
                    params.append(value)
                elif operator == 'not_equal':
                    condition = f"{table_field} != %s"
                    params.append(value)
                elif operator == 'contains':
                    condition = f"{table_field} LIKE %s"
                    params.append(f'%{value}%')
                elif operator == 'not_contains':
                    condition = f"{table_field} NOT LIKE %s"
                    params.append(f'%{value}%')
                elif operator == 'empty':
                    condition = f"({table_field} IS NULL OR {table_field} = '')"
                elif operator == 'not_empty':
                    condition = f"({table_field} IS NOT NULL AND {table_field} != '')"
            
            if condition:
                query_parts.append(condition)
        
        if not query_parts:
            return jsonify({'error': 'No valid search criteria provided'}), 400
        
        where_clause = f" {logic} ".join(query_parts)
        
        visible_columns = data.get('columns', 'p_id,p_status,p_apn,or_name,p_county,p_state,p_acres,p_comp_market_value').split(',')

        if 'or_name' in visible_columns or any(col.startswith('or_') or col.startswith('o_') for col in visible_columns):
            needs_owners = True
        
        page = int(data.get('page', 1))
        per_page = int(data.get('per_page', 10))
        sort_by = data.get('sort_by', 'p_id')
        sort_dir = data.get('sort_dir', 'asc')
        
        allowed_sort_fields = {'p_id', 'p_apn', 'p_county', 'p_state', 'p_price', 'p_acres'}
        if sort_by not in allowed_sort_fields:
            sort_by = 'p_id'
        
        sort_dir = 'DESC' if sort_dir.lower() == 'desc' else 'ASC'
        
        select_fields = ['p.p_id', 'p.or_id']
        
        for col in visible_columns:
            if col in ('p_id', 'or_id'):
                continue
            elif col == 'or_name':
                select_fields.append("CASE WHEN o.o_type = 'Company' THEN COALESCE(o.o_company, '') ELSE TRIM(COALESCE(o.or_fname, '') || ' ' || COALESCE(o.or_lname, '')) END as or_name")
                needs_owners = True
            elif col == 'tags':
                select_fields.append("(SELECT STRING_AGG(t2.tag_name, ' | ') FROM property_tags pt2 JOIN tags t2 ON pt2.tag_id = t2.tag_id WHERE pt2.p_id = p.p_id) as tags")
            elif col == 'p_status':
                select_fields.append('s.s_status as p_status')
                needs_statuses = True
            elif col.startswith('or_') or col.startswith('o_'):
                select_fields.append(f'o.{col}')
                needs_owners = True
            else:
                select_fields.append(f'p.{col}')
        
        join_clause_parts = []
        if needs_owners:
            join_clause_parts.append("JOIN owners o ON p.or_id = o.or_id")
        if needs_tags:
            join_clause_parts.append("LEFT JOIN property_tags pt ON p.p_id = pt.p_id LEFT JOIN tags t ON pt.tag_id = t.tag_id")
        if needs_statuses:
            join_clause_parts.append("LEFT JOIN statuses s ON p.p_status_id = s.status_id")

        join_clause = " ".join(join_clause_parts)
        
        # Count total
        count_sql = f"SELECT COUNT(DISTINCT p.p_id) FROM properties p {join_clause} WHERE {where_clause}"
        cursor.execute(count_sql, params)
        total_records = cursor.fetchone()['count']
        total_pages = (total_records + per_page - 1) // per_page
        
        fields_str = ', '.join(select_fields)
        query = f"""
            SELECT DISTINCT {fields_str}
            FROM properties p
            {join_clause}
            WHERE {where_clause}
            ORDER BY p.{sort_by} {sort_dir}
            LIMIT %s OFFSET %s
        """
        
        params.extend([per_page, (page - 1) * per_page])
        
        cursor.execute(query, params)
        rows = cursor.fetchall()
        
        properties = []
        for row in rows:
            properties.append(dict(row))
        
        return jsonify({
            'success': True,
            'properties': properties,
            'pagination': {
                'current_page': page,
                'per_page': per_page,
                'total_records': total_records,
                'total_pages': total_pages
            },
            'search_criteria': {
                'logic': logic,
                'rules_count': len(rules)
            }
        })
        
    except Exception as e:
        print(f"ERROR in advanced search: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()


                    
# ============================================
# THIS MUST BE THE LAST THING IN THE FILE
# ============================================
if __name__ == '__main__':
    import sys
    if '--dev' in sys.argv:
        # Only for debugging
        app.run(debug=True, port=8000)
    else:
        # Production mode (default)
        from waitress import serve
        print("Starting Valiant Land Server...")
        print("Open http://localhost:8000 in your browser")
        serve(app, host='0.0.0.0', port=8000, threads=4)